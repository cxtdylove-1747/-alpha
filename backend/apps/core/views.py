import os
import re
import tempfile
import threading
import zipfile
from difflib import SequenceMatcher
from io import BytesIO
from collections import defaultdict
from datetime import timedelta
from html import unescape
from pathlib import Path

try:
    from docx import Document
except Exception:  # pragma: no cover - optional import guard
    Document = None

from django.contrib.auth import get_user_model
from django.core.files import File
from django.db import close_old_connections
from django.db import transaction
from django.db.models import Count
from django.db.models import Q
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.utils.dateparse import parse_datetime
from rest_framework import generics, permissions, status
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework_simplejwt.tokens import RefreshToken

from ..agents.competition_advisor import CompetitionAdvisor
from ..agents.competition_coach import CompetitionCoachAgent
from ..agents.financial_designer import FinancialDesignAgent
from ..agents.langgraph_workflow import LangGraphWorkflowAgent
from ..agents.pitch_agents import PitchOptimizeAgent, PitchSimulationAgent
from ..agents.project_coach import ProjectCoach
from ..agents.roadmap_simulator import RoadmapSimulator
from ..agents.tutor_agent import TutorAgent
from .services.class_learning_report import build_class_learning_report
from .services.graph_visualization import build_hypergraph_preview, build_knowledge_graph_preview
from .services.graph_observability import track_graph_call
from .services.hypergraph_reasoner import infer_project_paths
from .services.hypergraph_runtime import build_plan_diagnosis, get_hypergraph_client
from .services.potential_report import build_potential_report
from .services.project_quality import evaluate_project_quality
from .services.vector_retrieval import VectorRetrievalService
from .services.hypergraph_preview import build_preview_graph
from .services.persona_profile import (
    build_class_persona,
    build_student_class_persona,
    build_student_persona,
    build_teacher_class_persona,
)
from ..hypergraph.neo4j_client import get_neo4j_client

from .ai_engine import (
    diagnose_pdf,
    next_guiding_question,
    recommend_knowledge,
    summarize_common_issues,
    summarize_common_issues_from_review_outputs,
    summarize_common_issues_from_plan_texts,
    student_pdf_chat_answer,
    teacher_chat_answer,
    validate_logic,
)
from .case_ingest import _run_office_to_pdf
from .models import (
    CaseLibraryDocument,
    CommonIssueReport,
    ConversationRecord,
    MentorshipApplication,
    MentorshipRelation,
    MessageNotification,
    Plan,
    ProjectRubricScore,
    PromptSceneConfig,
    EvidenceRecord,
    GraphInvocationLog,
    Rule,
    RubricItem,
    ReviewRecord,
    ScoringRubric,
    TeacherIntervention,
    TeacherChatRecord,
)
from .pdf_utils import parse_pdf
from .permissions import IsPlatformAdmin, IsStudent, IsTeacher
from .serializers import (
    CaseLibraryDocumentSerializer,
    CommonIssueReportSerializer,
    ConversationRecordSerializer,
    MentorshipApplicationSerializer,
    MessageNotificationSerializer,
    PlanSerializer,
    PromptSceneConfigSerializer,
    ProjectRubricScoreSerializer,
    RegisterSerializer,
    ReviewRecordSerializer,
    ScoringRubricSerializer,
    EvidenceRecordSerializer,
    TeacherInterventionSerializer,
    TeacherChatRecordSerializer,
    UserSerializer,
)
from .utils import log_event

User = get_user_model()

RUBRIC_NAME_MAP = {
    "R1": "Problem Definition",
    "R2": "User Evidence Strength",
    "R3": "Solution Feasibility",
    "R4": "Business Model Consistency",
    "R5": "Market & Competition",
    "R6": "Financial Logic",
    "R7": "Innovation & Differentiation",
    "R8": "Team & Execution",
    "R9": "Presentation & Material Quality",
}
SEVERITY_ORDER = {"high": 3, "medium": 2, "low": 1}
COMMON_ISSUE_CATEGORY_RULES = [
    {
        "key": "problem_definition",
        "label": "问题定义与场景聚焦不足",
        "summary": "常见表现是痛点描述偏泛、应用场景不够收敛，导致后续方案价值不够聚焦。",
        "priority": 1,
        "keywords": ["痛点", "问题", "场景", "背景", "定位", "切入点", "需求", "范围漂移", "problem"],
    },
    {
        "key": "user_evidence",
        "label": "用户研究与证据支撑不足",
        "summary": "常见表现是访谈、问卷、试点或数据来源不足，关键判断缺少可追溯证据。",
        "priority": 1,
        "keywords": ["用户", "访谈", "问卷", "样本", "调研", "证据", "验证", "数据来源", "来源", "引用", "可信度", "evidence"],
    },
    {
        "key": "solution_feasibility",
        "label": "方案可行性与落地论证不足",
        "summary": "常见表现是技术实现路径、资源约束、试点方式或 MVP 设计解释不充分。",
        "priority": 2,
        "keywords": ["技术", "实现", "可行", "前提", "限制", "资源", "落地", "mvp", "试点", "技术路线", "feasibility"],
    },
    {
        "key": "business_model",
        "label": "商业模式闭环不完整",
        "summary": "常见表现是客户、价值主张、获客渠道、收入来源与付费逻辑之间衔接不紧。",
        "priority": 1,
        "keywords": ["商业模式", "盈利", "收入", "成本", "客户", "渠道", "获客", "转化", "付费", "单价", "business"],
    },
    {
        "key": "market_competition",
        "label": "市场规模与竞品分析不足",
        "summary": "常见表现是市场边界过大、竞品对比粗略，难以证明项目切入空间与竞争优势。",
        "priority": 2,
        "keywords": ["市场", "竞品", "竞争", "tam", "sam", "som", "行业", "份额", "对比", "market", "competition"],
    },
    {
        "key": "financial_logic",
        "label": "财务测算与关键假设不足",
        "summary": "常见表现是收入预测、成本结构、现金流或单位经济缺少明确假设和计算依据。",
        "priority": 1,
        "keywords": ["财务", "预测", "现金流", "成本结构", "毛利", "利润", "测算", "单位经济", "ltv", "cac", "financial"],
    },
    {
        "key": "innovation_diff",
        "label": "创新点与差异化证明不足",
        "summary": "常见表现是创新主张停留在口号层，缺少量化对比和可验证指标。",
        "priority": 2,
        "keywords": ["创新", "差异化", "独特", "护城河", "优势", "壁垒", "量化对比", "innovation"],
    },
    {
        "key": "execution_team",
        "label": "执行计划与团队支撑不足",
        "summary": "常见表现是里程碑、资源配置、团队分工和交付节奏之间缺少对应关系。",
        "priority": 2,
        "keywords": ["团队", "执行", "里程碑", "计划", "分工", "交付", "时间表", "资源映射", "execution"],
    },
    {
        "key": "presentation_quality",
        "label": "材料结构与表达规范问题",
        "summary": "常见表现是标题、摘要、目录、图表、格式或叙事表达影响评委的阅读效率与专业感。",
        "priority": 3,
        "keywords": ["封面", "标题", "摘要", "目录", "图表", "页码", "格式", "语病", "重复", "表述", "专业性", "第一印象", "叙事", "presentation"],
    },
    {
        "key": "consistency_logic",
        "label": "前后逻辑一致性不足",
        "summary": "常见表现是章节之间存在矛盾、跳跃或概念混用，削弱整体说服力。",
        "priority": 1,
        "keywords": ["矛盾", "不一致", "冲突", "前后", "逻辑", "跳跃", "混淆", "断裂", "一致性", "consistency"],
    },
    {
        "key": "generic_argument",
        "label": "核心论证仍需加强",
        "summary": "常见表现是关键结论已经提出，但支撑论据、结构衔接或说服力仍不够稳定。",
        "priority": 4,
        "keywords": [],
    },
]


def _read_int_env(name: str, default: int) -> int:
    try:
        return int(os.getenv(name, str(default)) or str(default))
    except Exception:
        return default


def _normalize_issue_text(text: str) -> str:
    value = re.sub(r"\s+", " ", str(text or "").strip())
    return value[:220]


def _classify_common_issue(text: str) -> dict:
    normalized = str(text or "").strip().lower()
    best_rule = COMMON_ISSUE_CATEGORY_RULES[-1]
    best_score = -1
    for rule in COMMON_ISSUE_CATEGORY_RULES:
        keywords = rule.get("keywords") or []
        score = sum(1 for keyword in keywords if keyword and keyword.lower() in normalized)
        if score > best_score:
            best_score = score
            best_rule = rule
        elif score == best_score and score > 0 and int(rule.get("priority") or 99) < int(best_rule.get("priority") or 99):
            best_rule = rule
    return best_rule


def _build_common_issue_rows(complete_plans, review_by_plan):
    issue_groups = {}
    for plan in complete_plans:
        generated = review_by_plan.get(plan.id) or {}
        for issue in (generated.get("issues") or []):
            issue_payload = issue if isinstance(issue, dict) else {"description": issue}
            issue_text = _normalize_issue_text(issue_payload.get("description") or "")
            if not issue_text:
                continue
            category = _classify_common_issue(issue_text)
            key = str(category.get("key") or "generic_argument")
            row = issue_groups.setdefault(
                key,
                {
                    "issue": category.get("label") or "核心论证仍需加强",
                    "summary": category.get("summary") or "",
                    "risk": "low",
                    "priority": int(category.get("priority") or 99),
                    "mention_count": 0,
                    "project_ids": set(),
                    "projects": [],
                    "examples": [],
                },
            )
            row["mention_count"] += 1
            row["project_ids"].add(plan.id)
            risk = str(issue_payload.get("risk_level") or "medium").lower()
            if SEVERITY_ORDER.get(risk, 0) > SEVERITY_ORDER.get(row["risk"], 0):
                row["risk"] = risk
            if len(row["projects"]) < 8 and not any(item.get("project_id") == plan.id for item in row["projects"]):
                row["projects"].append(
                    {
                        "project_id": plan.id,
                        "project_name": plan.title,
                        "student_name": plan.student.display_name,
                    }
                )
            if issue_text not in row["examples"] and len(row["examples"]) < 2:
                row["examples"].append(issue_text)

    rows = [
        {
            "issue": row["issue"],
            "summary": row["summary"],
            "count": len(row["project_ids"]),
            "mention_count": row["mention_count"],
            "risk": row["risk"],
            "projects": row["projects"],
            "examples": row["examples"],
            "_priority": row["priority"],
        }
        for row in issue_groups.values()
    ]
    rows = sorted(
        rows,
        key=lambda item: (
            -int(item.get("count") or 0),
            -SEVERITY_ORDER.get(str(item.get("risk") or "low").lower(), 0),
            int(item.get("_priority") or 99),
            -int(item.get("mention_count") or 0),
        ),
    )
    repeated_rows = [item for item in rows if int(item.get("count") or 0) >= 2]
    target_rows = repeated_rows or rows
    return [{k: v for k, v in item.items() if k != "_priority"} for item in target_rows[:5]]


REVIEW_TASK_TIMEOUT_SECONDS = max(120, _read_int_env("REVIEW_TASK_TIMEOUT_SECONDS", 420))


def ok(data=None, message="ok", status_code=status.HTTP_200_OK):
    return Response({"code": status_code, "message": message, "data": data}, status=status_code)


def fail(message, status_code=status.HTTP_400_BAD_REQUEST):
    return Response({"code": status_code, "message": message, "data": None}, status=status_code)


def _plan_file_ext(plan: Plan) -> str:
    return Path(getattr(plan.pdf_file, "name", "") or "").suffix.lower()


def _build_pages_from_text(text: str, chunk_size: int = 2200):
    content = (text or "").strip()
    if not content:
        return []
    pages = []
    cursor = 0
    page_no = 1
    while cursor < len(content):
        chunk = content[cursor: cursor + chunk_size].strip()
        if chunk:
            pages.append({"page": page_no, "text": chunk})
            page_no += 1
        cursor += chunk_size
    return pages or [{"page": 1, "text": content}]


def _parse_non_pdf_document(file_path: str, file_ext: str):
    def _extract_docx_text(path: str) -> str:
        try:
            lines = []
            with zipfile.ZipFile(path, "r") as zf:
                xml_names = [
                    name
                    for name in zf.namelist()
                    if name in {
                        "word/document.xml",
                        "word/header1.xml",
                        "word/header2.xml",
                        "word/footer1.xml",
                        "word/footer2.xml",
                        "word/footnotes.xml",
                        "word/endnotes.xml",
                    }
                ]
                for name in xml_names:
                    raw = zf.read(name).decode("utf-8", errors="ignore")
                    raw = raw.replace("</w:p>", "\n").replace("<w:tab/>", "\t")
                    text = re.sub(r"<[^>]+>", "", raw)
                    text = unescape(text)
                    text = re.sub(r"\n{3,}", "\n\n", text)
                    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
                    if text:
                        lines.append(text)
            return "\n".join(lines).strip()
        except Exception:
            return ""

    text = ""
    if file_ext == ".docx" and Document is not None:
        try:
            doc = Document(file_path)
            lines = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
            text = "\n".join(lines).strip()
        except Exception:
            text = ""

    if not text and file_ext == ".docx":
        text = _extract_docx_text(file_path)

    if not text and file_ext == ".doc":
        # Try converting legacy .doc to PDF when possible; fallback to plain-text decode.
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                converted_pdf = _run_office_to_pdf(file_path, temp_dir)
                if converted_pdf:
                    return parse_pdf(converted_pdf)
        except Exception:
            pass

    if not text and file_ext == ".doc":
        try:
            raw = Path(file_path).read_bytes()
            text = raw.decode("utf-8", errors="ignore").strip()
            if not text:
                text = raw.decode("gb18030", errors="ignore").strip()
        except Exception:
            text = ""

    pages = _build_pages_from_text(text)
    return {
        "text": text,
        "pages": pages,
        "page_count": len(pages),
    }


def _parse_plan_file(file_path: str, file_ext: str):
    if file_ext == ".pdf":
        return parse_pdf(file_path)
    return _parse_non_pdf_document(file_path, file_ext)


def _ensure_plan_pages(plan: Plan):
    pages = plan.parsed_pages or []
    if any((p or {}).get("text", "").strip() for p in pages):
        return pages
    if not plan.pdf_file:
        return []

    ext = _plan_file_ext(plan)
    parsed = _parse_plan_file(plan.pdf_file.path, ext)
    parsed_pages = parsed.get("pages", []) or []
    parsed_text = parsed.get("text", "") or ""
    parsed_page_count = parsed.get("page_count", len(parsed_pages)) or len(parsed_pages)

    update_fields = []
    if parsed_pages:
        plan.parsed_pages = parsed_pages
        update_fields.append("parsed_pages")
    if parsed_text:
        plan.extracted_text = parsed_text
        update_fields.append("extracted_text")
    if parsed_page_count:
        plan.page_count = parsed_page_count
        update_fields.append("page_count")
    if update_fields:
        plan.save(update_fields=update_fields)

    return parsed_pages


def _to_project_group_key(plan: Plan) -> str:
    return f"{plan.student_id}:{(plan.title or '').strip().lower()}"


def _is_plan_complete(plan: Plan) -> bool:
    latest = ReviewRecord.objects.filter(plan=plan).order_by("-created_at").first()
    meta = (latest.review_meta if latest else {}) or {}
    if "is_complete" in meta:
        return bool(meta.get("is_complete"))
    text = (plan.extracted_text or "").strip()
    return len(text) >= 500 and all(k in text for k in ["鐢ㄦ埛", "鏂规", "鍟嗕笟"])


def _plan_progress(plan: Plan) -> int:
    latest = ReviewRecord.objects.filter(plan=plan).order_by("-created_at").first()
    meta = (latest.review_meta if latest else {}) or {}
    if "progress" in meta:
        try:
            return int(meta.get("progress") or 0)
        except Exception:
            return 0
    return 100 if _is_plan_complete(plan) else 50


def _delete_plan_with_local_file(plan: Plan):
    if plan.pdf_file:
        local_path = plan.pdf_file.path if hasattr(plan.pdf_file, "path") else ""
        try:
            plan.pdf_file.delete(save=False)
        except Exception:
            pass
        if local_path and os.path.exists(local_path):
            try:
                os.remove(local_path)
            except Exception:
                pass
    plan.delete()


def _build_hypergraph_meta(hg_report: dict) -> dict:
    report = hg_report or {}
    metrics = report.get("metrics") or {}
    consistency_alerts = report.get("consistency_alerts") or report.get("warnings") or []
    missing_labels = report.get("missing_node_labels") or report.get("missing_key_nodes") or []
    risk_patterns = report.get("risk_patterns") or []
    node_sources = (report.get("provenance") or {}).get("node_sources") or []
    edge_evidence = (report.get("provenance") or {}).get("hyperedge_evidence") or []
    risk_count = len(risk_patterns)
    alert_count = len(consistency_alerts)
    missing_count = len(missing_labels)

    def _readable_item_text(item) -> str:
        if item is None:
            return ""
        if isinstance(item, (str, int, float, bool)):
            return str(item).strip()
        if isinstance(item, dict):
            preferred_keys = ["text", "summary", "detail", "reason", "conclusion", "message", "name", "label"]
            parts = [str(item.get(key) or "").strip() for key in preferred_keys if str(item.get(key) or "").strip()]
            if parts:
                return "；".join(parts[:3])
            fallback = []
            for key in ("competition", "evidence_type", "file", "source_file", "path", "type"):
                value = str(item.get(key) or "").strip()
                if value:
                    fallback.append(f"{key}: {value}")
            return "；".join(fallback[:3])
        if isinstance(item, (list, tuple, set)):
            parts = [_readable_item_text(row) for row in item]
            return "；".join([part for part in parts if part][:3])
        return str(item).strip()

    if risk_count or alert_count or missing_count:
        conclusion = "当前项目存在结构性风险，建议先补齐关键证据与缺失节点，再进入下一轮扩展。"
    elif float(metrics.get("label_coverage_rate") or 0) >= 80 and float(metrics.get("explainability_item_rate") or 0) >= 60:
        conclusion = "当前超图结构较完整，可继续推进落地并保持证据链持续更新。"
    else:
        conclusion = "当前超图可用但完整度一般，建议继续增强关键标签覆盖和可解释性证据。"
    findings = [
        {
            "key": "overall",
            "title": "总体判断",
            "level": "warning" if (risk_count or alert_count or missing_count) else "positive",
            "detail": conclusion,
            "evidence": [
                f"标签覆盖率 {metrics.get('label_coverage_rate', 0)}%",
                f"可解释项覆盖 {metrics.get('explainability_item_rate', 0)}%",
                f"项目匹配置信度 {metrics.get('project_confidence_rate', 0)}%",
            ],
            "action": "先处理结构缺口，再继续放大方案规模。" if (risk_count or alert_count or missing_count) else "保持当前结构，继续补充高质量业务证据。",
        },
        {
            "key": "consistency",
            "title": "结构一致性",
            "level": "warning" if alert_count else "positive",
            "detail": f"发现 {alert_count} 条一致性告警。" if alert_count else "暂未发现明显结构冲突。",
            "evidence": consistency_alerts[:6],
            "action": "优先修正告警对应的事实链和逻辑链。" if alert_count else "继续保持当前结构一致性。",
        },
        {
            "key": "missing_nodes",
            "title": "关键节点覆盖",
            "level": "warning" if missing_count else "positive",
            "detail": f"存在 {missing_count} 项缺失关键节点。" if missing_count else "关键节点覆盖较完整。",
            "evidence": missing_labels[:8],
            "action": (
                f"建议优先补齐：{', '.join([_readable_item_text(item) for item in missing_labels[:4] if _readable_item_text(item)])}"
                if missing_count
                else "后续重点提升节点之间的证据关联强度。"
            ),
        },
        {
            "key": "risk_patterns",
            "title": "风险模式识别",
            "level": "warning" if risk_count else "neutral",
            "detail": f"识别到 {risk_count} 条风险模式。" if risk_count else "暂未识别到明显高风险模式。",
            "evidence": [
                f"{_readable_item_text(item if not isinstance(item, dict) else (item.get('type') or '风险模式'))}（{str(item.get('severity') or '中') if isinstance(item, dict) else '中'}）"
                for item in risk_patterns[:8]
            ],
            "action": "对高频风险模式补证据、补边界条件、补量化指标。" if risk_count else "继续观察关键风险触发点。",
        },
        {
            "key": "provenance",
            "title": "证据可追溯性",
            "level": "positive" if (len(node_sources) + len(edge_evidence)) else "warning",
            "detail": f"当前共关联 {len(node_sources) + len(edge_evidence)} 条溯源证据。",
            "evidence": [
                f"节点溯源 {len(node_sources)} 条",
                f"超边溯源 {len(edge_evidence)} 条",
                f"溯源命中率 {metrics.get('source_hit_rate', 0)}%",
            ],
            "action": "继续补充关键结论对应的原始片段和来源文件。",
        },
    ]
    if report.get("similar_projects"):
        findings.append(
            {
                "key": "similar_projects",
                "title": "相似项目参照",
                "level": "neutral",
                "detail": f"匹配到 {len(report.get('similar_projects') or [])} 个相似项目，可用于校准结构和证据。",
                "evidence": [
                    (
                        str(item.get("name") or item.get("project_name") or item.get("id") or "相似项目")
                        if isinstance(item, dict)
                        else _readable_item_text(item) or "相似项目"
                    )
                    for item in (report.get("similar_projects") or [])[:6]
                ],
                "action": "对比相似项目的市场定义、指标设计和论证链完整度。",
            }
        )
    if report.get("suggested_evidence"):
        findings.append(
            {
                "key": "suggested_evidence",
                "title": "建议补证方向",
                "level": "warning",
                "detail": f"当前建议补充 {len(report.get('suggested_evidence') or [])} 条证据。",
                "evidence": [
                    _readable_item_text(item)[:120]
                    for item in (report.get("suggested_evidence") or [])[:6]
                    if _readable_item_text(item)
                ],
                "action": "优先补齐能直接支撑市场、指标、风险控制的证据。",
            }
        )
    summary_evidence = [
        f"标签覆盖率 {metrics.get('label_coverage_rate', 0)}%，可解释项覆盖 {metrics.get('explainability_item_rate', 0)}%。",
        f"一致性告警 {alert_count} 条，缺失节点 {missing_count} 项，风险模式 {risk_count} 条。",
        f"溯源证据 {len(node_sources) + len(edge_evidence)} 条（节点 {len(node_sources)}，超边 {len(edge_evidence)}）。",
    ]
    return {
        "enabled": bool(report.get("enabled")),
        "input_type": report.get("input_type"),
        "matched_project": report.get("matched_project"),
        "matched_project_node_id": report.get("matched_project_node_id"),
        "matched_project_name": report.get("matched_project_name"),
        "resolved_project_node_id": report.get("resolved_project_node_id"),
        "consistency_alerts": consistency_alerts,
        "missing_node_labels": missing_labels,
        "similar_projects": report.get("similar_projects") or [],
        "risk_patterns": risk_patterns,
        "suggested_evidence": report.get("suggested_evidence") or [],
        "entity_match_stats": report.get("entity_match_stats") or {},
        "metrics": metrics,
        "diagnosis_outline": report.get("diagnosis_outline") or [],
        "provenance": report.get("provenance") or {"node_sources": [], "hyperedge_evidence": []},
        "diagnosis_summary": {
            "conclusion": conclusion,
            "evidence": summary_evidence,
            "findings": findings,
        },
        # Keep compatibility fields for existing consumers.
        "warnings": consistency_alerts,
        "missing_key_nodes": missing_labels,
    }


def _build_hypergraph_suggestions(hg_report: dict) -> list[str]:
    report = hg_report or {}
    suggestions = []
    alerts = (report.get("consistency_alerts") or report.get("warnings") or [])[:3]
    for item in alerts:
        suggestions.append(f"[超图校验] {item}")
    missing = report.get("missing_node_labels") or report.get("missing_key_nodes") or []
    if missing:
        suggestions.append(f"[超图补齐] 建议优先补充：{', '.join(missing[:4])}")
    metrics = report.get("metrics") or {}
    if metrics:
        suggestions.append(
            "[超图指标] 关键标签覆盖率"
            f"{metrics.get('label_coverage_rate', 0)}%，可解释项覆盖"
            f"{metrics.get('explainability_item_rate', 0)}%。"
        )
    return suggestions


def _dedupe_text_rows(values) -> list[str]:
    seen = set()
    rows = []
    for item in values or []:
        text = str(item or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        rows.append(text)
    return rows


def _normalize_rubric_dimension_id(value) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    upper = raw.upper()
    if upper in RUBRIC_NAME_MAP:
        return upper

    normalized = re.sub(r"[^a-z0-9\u4e00-\u9fa5]+", " ", raw.replace("&", " and ")).strip().lower()
    label_aliases = {
        "problem definition": "R1",
        "user evidence strength": "R2",
        "solution feasibility": "R3",
        "business model consistency": "R4",
        "market competition": "R5",
        "market and competition": "R5",
        "financial logic": "R6",
        "innovation differentiation": "R7",
        "innovation and differentiation": "R7",
        "team execution": "R8",
        "team and execution": "R8",
        "presentation material quality": "R9",
        "presentation and material quality": "R9",
    }
    return label_aliases.get(normalized, "")


def _normalize_review_dimension_scores(rows) -> list[dict]:
    merged = {}
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        rubric_id = _normalize_rubric_dimension_id(
            row.get("code") or row.get("label") or row.get("dimension") or row.get("name")
        )
        if not rubric_id:
            continue
        try:
            score = round(float(row.get("score") or 0), 2)
        except Exception:
            score = 0.0
        normalized = dict(row)
        normalized["code"] = rubric_id
        normalized["label"] = RUBRIC_NAME_MAP.get(rubric_id, rubric_id)
        normalized["score"] = score
        merged[rubric_id] = normalized
    ordered = []
    for idx in range(1, len(RUBRIC_NAME_MAP) + 1):
        rubric_id = f"R{idx}"
        if rubric_id in merged:
            ordered.append(merged[rubric_id])
    return ordered


def _review_meta_has_current_rubric(review_meta: dict) -> bool:
    if not isinstance(review_meta, dict):
        return False
    dimension_scores = _normalize_review_dimension_scores(review_meta.get("dimension_scores") or [])
    return len(dimension_scores) == len(RUBRIC_NAME_MAP)


def _build_review_payload_dict(record: ReviewRecord, review_meta: dict) -> dict:
    normalized_dimension_scores = _normalize_review_dimension_scores(review_meta.get("dimension_scores") or [])
    return {
        "id": record.id,
        "issues": record.issues or [],
        "annotations": record.annotations or [],
        "guidance_questions": record.guidance_questions or [],
        "examples": record.examples or [],
        "suggestions": record.suggestions or [],
        "review_meta": review_meta,
        "summary": record.summary or "",
        "dimension_scores": normalized_dimension_scores,
    }


def _generate_review_payload(plan: Plan, audience_role: str) -> dict:
    pages = _ensure_plan_pages(plan)
    if not pages:
        pages = [{"page": 1, "text": plan.extracted_text or plan.title or ""}]
    full_text = "\n".join((p or {}).get("text", "") for p in pages).strip()
    project_type = _normalize_project_type("auto", plan_text=full_text)
    hg_report = build_plan_diagnosis(plan)
    payload = diagnose_pdf(
        pages,
        audience_role,
        hypergraph_context=hg_report,
        project_type=project_type,
    )
    payload.setdefault("review_meta", {})
    payload["review_meta"]["parsed_page_count"] = len(pages)
    payload["review_meta"]["parsed_text_length"] = len(full_text)
    payload["review_meta"]["project_type"] = project_type
    payload["review_meta"]["hypergraph"] = _build_hypergraph_meta(hg_report)

    normalized_dimension_scores = _normalize_review_dimension_scores(
        (payload.get("review_meta") or {}).get("dimension_scores") or payload.get("dimension_scores") or []
    )
    payload["review_meta"]["dimension_scores"] = normalized_dimension_scores
    payload["dimension_scores"] = normalized_dimension_scores

    if hg_report.get("enabled"):
        hg_suggestions = _build_hypergraph_suggestions(hg_report)
        if hg_suggestions:
            payload["suggestions"] = _dedupe_text_rows((payload.get("suggestions") or []) + hg_suggestions)
    else:
        payload["suggestions"] = _dedupe_text_rows(payload.get("suggestions") or [])
    return payload


def _get_or_create_review_payload(plan: Plan, audience_role: str, reviewer: User | None = None) -> dict:
    latest = ReviewRecord.objects.filter(plan=plan, audience_role=audience_role).order_by("-created_at").first()
    if latest:
        review_meta = (latest.review_meta or {}).copy()
        normalized_dimension_scores = _normalize_review_dimension_scores(review_meta.get("dimension_scores") or [])
        needs_rubric_backfill = not _review_meta_has_current_rubric(review_meta)
        if needs_rubric_backfill:
            payload = _generate_review_payload(plan, audience_role)
            review_meta = (payload.get("review_meta") or {}).copy()
            latest.reviewer = reviewer or latest.reviewer
            latest.issues = payload.get("issues") or []
            latest.annotations = payload.get("annotations") or []
            latest.guidance_questions = payload.get("guidance_questions") or []
            latest.examples = payload.get("examples") or []
            latest.suggestions = payload.get("suggestions") or []
            latest.review_meta = review_meta
            latest.summary = payload.get("summary") or ""
            latest.save(
                update_fields=[
                    "reviewer",
                    "issues",
                    "annotations",
                    "guidance_questions",
                    "examples",
                    "suggestions",
                    "review_meta",
                    "summary",
                ]
            )
            return _build_review_payload_dict(latest, review_meta)

        if not isinstance(review_meta.get("hypergraph"), dict) or review_meta.get("dimension_scores") != normalized_dimension_scores:
            hg_report = build_plan_diagnosis(plan)
            review_meta["hypergraph"] = _build_hypergraph_meta(hg_report)
            review_meta["dimension_scores"] = normalized_dimension_scores
            latest.review_meta = review_meta
            latest.save(update_fields=["review_meta"])
        return _build_review_payload_dict(latest, review_meta)

    payload = _generate_review_payload(plan, audience_role)

    record = ReviewRecord.objects.create(
        plan=plan,
        reviewer=reviewer,
        audience_role=audience_role,
        issues=payload.get("issues") or [],
        annotations=payload.get("annotations") or [],
        guidance_questions=payload.get("guidance_questions") or [],
        examples=payload.get("examples") or [],
        suggestions=payload.get("suggestions") or [],
        review_meta=payload.get("review_meta") or {},
        summary=payload.get("summary") or "",
    )
    return _build_review_payload_dict(record, payload.get("review_meta") or {})


class HealthView(APIView):
    permission_classes = [permissions.AllowAny]

    def get(self, request):
        retrieval_mode = VectorRetrievalService().retrieval_mode()
        workflow_mode = LangGraphWorkflowAgent.mode()
        return ok(
            {
                "status": "ok",
                "service": "innovation-agent-backend",
                "retrieval_mode": retrieval_mode,
                "workflow_mode": workflow_mode,
                "time": timezone.now().isoformat(),
            }
        )


class RegisterView(generics.CreateAPIView):
    permission_classes = [permissions.AllowAny]
    serializer_class = RegisterSerializer


class LoginView(APIView):
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        username = str(request.data.get("username", "") or "").strip()
        password = request.data.get("password", "")
        user = User.objects.filter(
            Q(username=username) | Q(student_no=username) | Q(teacher_no=username)
        ).first()
        if not user or not user.check_password(password):
            return fail("账号或密码错误")
        if not user.is_active:
            return fail("账号已停用，请联系管理员")

        refresh = RefreshToken.for_user(user)
        return ok({"access": str(refresh.access_token), "refresh": str(refresh), "user": UserSerializer(user).data})


class LogoutView(APIView):
    def post(self, request):
        return ok(message="退出成功")


class ProfileView(APIView):
    def get(self, request):
        return ok(UserSerializer(request.user).data)

    def put(self, request):
        serializer = UserSerializer(instance=request.user, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return ok(serializer.data, message="更新成功")


class StudentPersonaView(APIView):
    permission_classes = [IsStudent]

    def get(self, request):
        return ok(build_student_persona(request.user))


class StudentClassPersonaView(APIView):
    permission_classes = [IsStudent]

    def get(self, request):
        return ok(build_student_class_persona(request.user))


class TeacherClassPersonaView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        student_ids = _teacher_student_ids(request.user, include_pending=False, allow_plan_fallback=True)
        return ok(build_class_persona(student_ids=student_ids, class_name=f"{request.user.username} 指导班级画像"))


class AdminOverviewView(APIView):
    permission_classes = [IsPlatformAdmin]

    def get(self, request):
        now = timezone.now()
        one_day_ago = now - timedelta(days=1)

        user_counts = User.objects.values("role").annotate(total=Count("id"))
        student_count = sum(row.get("total", 0) for row in user_counts if row.get("role") == User.ROLE_STUDENT)
        teacher_count = sum(row.get("total", 0) for row in user_counts if row.get("role") == User.ROLE_TEACHER)
        admin_count = User.objects.filter(Q(is_staff=True) | Q(is_superuser=True) | Q(role=User.ROLE_ADMINISTER)).count()

        plans_total = Plan.objects.count()
        plans_submitted = Plan.objects.filter(status=Plan.STATUS_SUBMITTED).count()
        plans_draft = Plan.objects.filter(status=Plan.STATUS_DRAFT).count()

        reviews_total = ReviewRecord.objects.count()
        chats_total = ConversationRecord.objects.count() + TeacherChatRecord.objects.count()
        mentorship_pending = MentorshipApplication.objects.filter(status=MentorshipApplication.STATUS_PENDING).count()

        recent_plans = list(
            Plan.objects.select_related("student").order_by("-created_at")[:6].values(
                "id", "title", "created_at", "student__username"
            )
        )
        recent_reviews = list(
            ReviewRecord.objects.select_related("plan").order_by("-created_at")[:6].values(
                "id", "created_at", "plan__title", "audience_role"
            )
        )

        activities = []
        for row in recent_plans:
            activities.append(
                {
                    "type": "plan",
                    "time": row.get("created_at"),
                    "title": f"新方案提交：{row.get('title') or '-'}",
                    "meta": f"学生：{row.get('student__username') or '-'}",
                }
            )
        for row in recent_reviews:
            activities.append(
                {
                    "type": "review",
                    "time": row.get("created_at"),
                    "title": f"新审阅生成：{row.get('plan__title') or '-'}",
                    "meta": f"面向角色：{row.get('audience_role') or '-'}",
                }
            )

        activities = sorted(
            activities,
            key=lambda x: (x.get("time").timestamp() if getattr(x.get("time"), "timestamp", None) else 0),
            reverse=True,
        )[:10]

        return ok(
            {
                "users": {
                    "total": User.objects.count(),
                    "students": student_count,
                    "teachers": teacher_count,
                    "admins": admin_count,
                    "new_24h": User.objects.filter(date_joined__gte=one_day_ago).count(),
                },
                "projects": {
                    "total": plans_total,
                    "submitted": plans_submitted,
                    "draft": plans_draft,
                    "new_24h": Plan.objects.filter(created_at__gte=one_day_ago).count(),
                },
                "reviews": {
                    "total": reviews_total,
                    "new_24h": ReviewRecord.objects.filter(created_at__gte=one_day_ago).count(),
                },
                "runtime": {
                    "retrieval_mode": getattr(VectorRetrievalService, "mode", "keyword"),
                    "workflow_mode": LangGraphWorkflowAgent.mode(),
                    "chat_records": chats_total,
                    "pending_mentorship": mentorship_pending,
                    "time": now.isoformat(),
                },
                "recent_activities": activities,
            }
        )


class AdminUserManageView(APIView):
    permission_classes = [IsPlatformAdmin]

    def get(self, request):
        keyword = str(request.query_params.get("q") or "").strip()
        role = str(request.query_params.get("role") or "all").strip().lower()
        is_active_raw = str(request.query_params.get("is_active") or "all").strip().lower()

        queryset = User.objects.all().order_by("-date_joined")
        if keyword:
            queryset = queryset.filter(
                Q(username__icontains=keyword)
                | Q(full_name__icontains=keyword)
                | Q(student_no__icontains=keyword)
                | Q(teacher_no__icontains=keyword)
                | Q(email__icontains=keyword)
                | Q(organization__icontains=keyword)
            )
        if role in {User.ROLE_STUDENT, User.ROLE_TEACHER, User.ROLE_ADMINISTER}:
            queryset = queryset.filter(role=role)
        if is_active_raw in {"true", "false"}:
            queryset = queryset.filter(is_active=(is_active_raw == "true"))

        rows = UserSerializer(queryset[:200], many=True).data
        return ok(rows)

    def post(self, request):
        overwrite = bool(request.data.get("overwrite_password"))
        default_role = str(request.data.get("default_role") or User.ROLE_STUDENT).strip().lower()
        if default_role not in {User.ROLE_STUDENT, User.ROLE_TEACHER}:
            default_role = User.ROLE_STUDENT

        users_payload = request.data.get("users")
        if isinstance(users_payload, list):
            rows = []
            for item in users_payload:
                if not isinstance(item, dict):
                    continue
                role = str(item.get("role") or default_role).strip().lower()
                if role not in {User.ROLE_STUDENT, User.ROLE_TEACHER}:
                    role = default_role

                account_no = str(
                    item.get("account_no")
                    or item.get("student_id")
                    or item.get("teacher_id")
                    or item.get("username")
                    or ""
                ).strip()
                password = str(item.get("password") or item.get("default_password") or "").strip()
                if not account_no or not password:
                    continue
                rows.append(
                    {
                        "username": account_no,
                        "account_no": account_no,
                        "password": password,
                        "role": role,
                        "full_name": str(item.get("name") or item.get("full_name") or "").strip(),
                        "email": str(item.get("email") or "").strip(),
                    }
                )
        else:
            rows = _parse_batch_users(request.data.get("raw") or "", default_role=default_role)

        if not rows:
            return fail("未读取到可导入用户，请使用：角色,姓名,学号/工号,默认密码,邮箱（每行一条）")

        created = []
        updated = []
        skipped = []
        errors = []

        with transaction.atomic():
            for row in rows:
                username = row.get("username")
                password = row.get("password")
                role = row.get("role") or default_role
                email = row.get("email") or ""
                full_name = row.get("full_name") or ""
                account_no = row.get("account_no") or username
                try:
                    user = (
                        User.objects.filter(username=username).first()
                        or User.objects.filter(student_no=account_no).first()
                        or User.objects.filter(teacher_no=account_no).first()
                    )
                    if not user:
                        user = User(username=username, role=role, email=email, is_active=True)
                        user.full_name = full_name
                        user.student_no = account_no if role == User.ROLE_STUDENT else None
                        user.teacher_no = account_no if role == User.ROLE_TEACHER else None
                        user.set_password(password)
                        user.save()
                        created.append(username)
                        continue

                    changed_fields = []
                    if user.username != username:
                        user.username = username
                        changed_fields.append("username")
                    if user.role != role and not (user.is_staff or user.is_superuser):
                        user.role = role
                        changed_fields.append("role")
                    if email and user.email != email:
                        user.email = email
                        changed_fields.append("email")
                    if full_name and user.full_name != full_name:
                        user.full_name = full_name
                        changed_fields.append("full_name")
                    next_student_no = account_no if role == User.ROLE_STUDENT else None
                    next_teacher_no = account_no if role == User.ROLE_TEACHER else None
                    if user.student_no != next_student_no:
                        user.student_no = next_student_no
                        changed_fields.append("student_no")
                    if user.teacher_no != next_teacher_no:
                        user.teacher_no = next_teacher_no
                        changed_fields.append("teacher_no")
                    if overwrite:
                        user.set_password(password)
                        changed_fields.append("password")
                    if changed_fields:
                        user.save()
                        updated.append(username)
                    else:
                        skipped.append(username)
                except Exception as exc:
                    errors.append({"username": username, "error": str(exc)})

        return ok(
            {
                "created": created,
                "updated": updated,
                "skipped": skipped,
                "errors": errors,
                "summary": {
                    "input_count": len(rows),
                    "created_count": len(created),
                    "updated_count": len(updated),
                    "skipped_count": len(skipped),
                    "error_count": len(errors),
                },
            },
            message="批量导入完成",
        )


class AdminUserStatusUpdateView(APIView):
    permission_classes = [IsPlatformAdmin]

    def post(self, request, user_id):
        target = get_object_or_404(User, id=user_id)
        if target.id == request.user.id:
            return fail("不能修改当前登录管理员自身状态")

        next_active = request.data.get("is_active")
        next_role = str(request.data.get("role") or "").strip().lower()
        next_staff = request.data.get("is_staff")
        next_superuser = request.data.get("is_superuser")

        changed_fields = []
        if isinstance(next_active, bool):
            target.is_active = next_active
            changed_fields.append("is_active")

        if isinstance(next_staff, bool):
            target.is_staff = next_staff
            changed_fields.append("is_staff")

        if isinstance(next_superuser, bool):
            target.is_superuser = next_superuser
            changed_fields.append("is_superuser")

        if next_role in {User.ROLE_STUDENT, User.ROLE_TEACHER, User.ROLE_ADMINISTER}:
            target.role = next_role
            changed_fields.append("role")

        # Keep role consistent with elevated platform privileges.
        if (target.is_staff or target.is_superuser) and target.role != User.ROLE_ADMINISTER:
            target.role = User.ROLE_ADMINISTER
            changed_fields.append("role")

        if not (target.is_staff or target.is_superuser) and target.role == User.ROLE_ADMINISTER:
            target.role = User.ROLE_TEACHER
            changed_fields.append("role")

        if not changed_fields:
            return fail("未提供有效更新字段")

        target.save(update_fields=sorted(set(changed_fields)))
        return ok(UserSerializer(target).data, message="用户状态已更新")


class AdminHypergraphGraphView(APIView):
    permission_classes = [IsPlatformAdmin]

    def get(self, request):
        # Prefer explicit hypergraph params; keep `limit` as backward-compatible fallback.
        limit_hyperedges = int(request.query_params.get("limit_hyperedges") or request.query_params.get("limit") or 120)
        max_members = int(request.query_params.get("max_members") or 6)
        payload = build_hypergraph_preview(limit_hyperedges=limit_hyperedges, max_members=max_members)
        stats = payload.get("stats") or {}
        load_error = str(stats.get("load_error") or "").strip()
        has_data = bool(payload.get("nodes"))
        if not has_data:
            return fail(load_error or "超图数据为空，请检查图谱构建文件", status.HTTP_503_SERVICE_UNAVAILABLE)
        return ok(payload)


class AdminKnowledgeGraphView(APIView):
    permission_classes = [IsPlatformAdmin]

    def get(self, request):
        limit_nodes = int(request.query_params.get("limit_nodes") or 220)
        limit_edges = int(request.query_params.get("limit_edges") or 500)
        payload = build_knowledge_graph_preview(limit_nodes=limit_nodes, limit_edges=limit_edges)
        stats = payload.get("stats") or {}
        load_error = str(stats.get("load_error") or "").strip()
        has_data = bool(payload.get("nodes"))
        if not has_data:
            return fail(load_error or "知识图谱数据为空，请检查图谱构建文件", status.HTTP_503_SERVICE_UNAVAILABLE)
        return ok(payload)


def _text_hit(left: str, right: str) -> bool:
    a = str(left or "").strip().lower()
    b = str(right or "").strip().lower()
    if not a or not b or len(b) < 3:
        return False
    return a in b or b in a


def _safe_ratio(numerator: int, denominator: int) -> float:
    if denominator <= 0:
        return 0.0
    return round((float(numerator) / float(denominator)) * 100, 2)


def _normalize_similarity_text(text: str) -> str:
    value = str(text or "").lower()
    value = re.sub(r"\s+", "", value)
    return value[:24000]


def _build_text_signature(text: str) -> set[str]:
    normalized = _normalize_similarity_text(text)
    chinese_chars = "".join(re.findall(r"[\u4e00-\u9fff]", normalized))
    cn_bigrams = {chinese_chars[idx: idx + 2] for idx in range(max(len(chinese_chars) - 1, 0))}
    words = set(re.findall(r"[a-z0-9]{3,}", normalized))
    signature = cn_bigrams | words
    if len(signature) > 2400:
        return set(sorted(signature)[:2400])
    return signature


def _jaccard_similarity(left: set[str], right: set[str]) -> float:
    if not left or not right:
        return 0.0
    union = left | right
    if not union:
        return 0.0
    return len(left & right) / len(union)


def _resolve_plan_text_for_check(plan: Plan) -> str:
    text = str(plan.extracted_text or "").strip()
    if text:
        return text
    try:
        pages = _ensure_plan_pages(plan)
    except Exception:
        pages = []
    return "\n".join((item or {}).get("text", "") for item in pages).strip()


def _build_plagiarism_check(selected_plans: list[Plan]) -> dict:
    entries = []
    for plan in selected_plans:
        text = _resolve_plan_text_for_check(plan)
        if len(text) < 120:
            continue
        entries.append(
            {
                "plan": plan,
                "text": text,
                "signature": _build_text_signature(text),
            }
        )

    suspicious_pairs = []
    for idx in range(len(entries)):
        left = entries[idx]
        for jdx in range(idx + 1, len(entries)):
            right = entries[jdx]
            score = _jaccard_similarity(left["signature"], right["signature"])
            if score <= 0:
                compact_left = _normalize_similarity_text(left["text"])[:3600]
                compact_right = _normalize_similarity_text(right["text"])[:3600]
                if compact_left and compact_right:
                    score = SequenceMatcher(None, compact_left, compact_right).ratio()

            if score < 0.42:
                continue

            risk_level = "high" if score >= 0.62 else "medium"
            suspicious_pairs.append(
                {
                    "plan_a_id": left["plan"].id,
                    "plan_a_name": left["plan"].title,
                    "student_a_name": left["plan"].student.display_name,
                    "plan_b_id": right["plan"].id,
                    "plan_b_name": right["plan"].title,
                    "student_b_name": right["plan"].student.display_name,
                    "similarity_score": round(score * 100, 1),
                    "risk_level": risk_level,
                    "reason": "文本结构与关键表述相似度偏高，建议进行剽窃复核。",
                }
            )

    suspicious_pairs = sorted(suspicious_pairs, key=lambda row: float(row.get("similarity_score") or 0), reverse=True)
    high_risk_count = sum(1 for item in suspicious_pairs if str(item.get("risk_level") or "").lower() == "high")
    return {
        "enabled": True,
        "checked_plan_count": len(entries),
        "suspicious_pair_count": len(suspicious_pairs),
        "high_risk_count": high_risk_count,
        "suspicious_pairs": suspicious_pairs[:20],
    }


def _safe_review_score_and_progress(plan: Plan, review_by_plan: dict, reviewer: User) -> tuple[float, int]:
    payload = review_by_plan.get(plan.id)
    if payload is None:
        payload = _get_or_create_review_payload(plan, ReviewRecord.ROLE_STUDENT, reviewer=reviewer)
        review_by_plan[plan.id] = payload

    review_meta = (payload.get("review_meta") or {})
    dim_rows = _normalize_review_dimension_scores(payload.get("dimension_scores") or review_meta.get("dimension_scores") or [])
    avg_score = round(sum(float(item.get("score") or 0) for item in dim_rows) / max(len(dim_rows), 1), 2) if dim_rows else 0.0
    progress = int(review_meta.get("progress") or _plan_progress(plan) or 0)
    return avg_score, progress


def _build_process_evaluation(selected_plans: list[Plan], grouped_versions: dict, review_by_plan: dict, reviewer: User) -> dict:
    rows = []
    ninety_days_ago = timezone.now() - timedelta(days=90)
    for plan in selected_plans:
        project_key = _to_project_group_key(plan)
        versions = sorted(grouped_versions.get(project_key) or [plan], key=lambda item: (item.version, item.created_at))
        version_ids = [int(item.id) for item in versions if getattr(item, "id", None)]

        chat_qs = ConversationRecord.objects.filter(user=plan.student)
        activity_count = (
            chat_qs.filter(Q(metadata__plan_id__in=version_ids) | Q(metadata__plan_id=int(plan.id))).count()
            if version_ids
            else 0
        )
        if activity_count <= 0:
            activity_count = chat_qs.filter(created_at__gte=ninety_days_ago).count()

        first_plan = versions[0]
        last_plan = versions[-1]
        first_score, first_progress = _safe_review_score_and_progress(first_plan, review_by_plan, reviewer)
        last_score, last_progress = _safe_review_score_and_progress(last_plan, review_by_plan, reviewer)
        score_delta = round(last_score - first_score, 2)
        progress_delta = int(last_progress - first_progress)
        iteration_count = len(versions)

        activity_score = min(100.0, float(activity_count) * 6.0)
        iteration_score = min(100.0, float(iteration_count) * 24.0)
        logic_score = max(0.0, min(100.0, 50.0 + float(score_delta) * 18.0 + float(progress_delta) * 0.6))
        process_score = round(activity_score * 0.35 + iteration_score * 0.25 + logic_score * 0.40, 1)

        level = "high" if process_score >= 75 else ("medium" if process_score >= 55 else "low")
        if level == "high":
            comment = "交互活跃且迭代充分，逻辑演进趋势明显。"
        elif level == "medium":
            comment = "过程参与度中等，建议增加高质量追问与结构化迭代。"
        else:
            comment = "过程证据偏弱，建议提升人机交互频次并记录版本迭代依据。"

        rows.append(
            {
                "project_id": plan.id,
                "project_name": plan.title,
                "student_name": plan.student.display_name,
                "activity_count": int(activity_count),
                "iteration_count": int(iteration_count),
                "first_score": round(first_score, 2),
                "latest_score": round(last_score, 2),
                "score_delta": round(score_delta, 2),
                "progress_delta": int(progress_delta),
                "process_score": process_score,
                "level": level,
                "comment": comment,
            }
        )

    rows = sorted(rows, key=lambda item: (float(item.get("process_score") or 0), -int(item.get("iteration_count") or 0)))
    low_count = sum(1 for item in rows if str(item.get("level") or "") == "low")
    high_count = sum(1 for item in rows if str(item.get("level") or "") == "high")
    avg_score = round(sum(float(item.get("process_score") or 0) for item in rows) / len(rows), 2) if rows else 0.0
    return {
        "overview": {
            "avg_process_score": avg_score,
            "project_count": len(rows),
            "high_process_count": high_count,
            "low_process_count": low_count,
        },
        "project_rows": rows[:24],
    }


def _build_class_warning_payload(
    selected_plans: list[Plan],
    incomplete_plans: list[Plan],
    review_by_plan: dict,
    plagiarism_check: dict,
    process_evaluation: dict,
) -> dict:
    total = len(selected_plans)
    incomplete_count = len(incomplete_plans)
    incomplete_rate = _safe_ratio(incomplete_count, total) if total else 0.0
    process_overview = process_evaluation.get("overview") or {}
    low_process_count = int(process_overview.get("low_process_count") or 0)
    high_issue_count = 0
    for plan in selected_plans:
        issues = (review_by_plan.get(plan.id) or {}).get("issues") or []
        if any(str((item or {}).get("risk_level") or "").lower() == "high" for item in issues if isinstance(item, dict)):
            high_issue_count += 1

    warnings = []
    plagiarism_high = int(plagiarism_check.get("high_risk_count") or 0)
    plagiarism_pairs = int(plagiarism_check.get("suspicious_pair_count") or 0)
    if plagiarism_high > 0:
        warnings.append(
            {
                "level": "high",
                "title": "疑似高风险剽窃",
                "detail": f"检测到 {plagiarism_high} 组高风险相似文本，请优先人工复核。",
                "action": "核查版本提交时间、访谈记录与原始素材证据链。",
            }
        )
    elif plagiarism_pairs > 0:
        warnings.append(
            {
                "level": "medium",
                "title": "存在中度文本相似",
                "detail": f"检测到 {plagiarism_pairs} 组疑似相似文本，建议抽样排查。",
                "action": "抽查项目原始调研记录与迭代日志，确认原创性。",
            }
        )

    if high_issue_count > 0:
        warnings.append(
            {
                "level": "high" if _safe_ratio(high_issue_count, total or 1) >= 30 else "medium",
                "title": "高风险逻辑问题偏多",
                "detail": f"{high_issue_count} 个项目存在高风险问题条目。",
                "action": "安排针对性答辩演练，优先修复高风险规则触发项。",
            }
        )

    if incomplete_rate >= 35:
        warnings.append(
            {
                "level": "high",
                "title": "班级项目完整度偏低",
                "detail": f"当前不完整项目占比 {incomplete_rate}%。",
                "action": "设置分阶段提交清单，优先补齐问题定义、证据与商业闭环。",
            }
        )
    elif incomplete_rate >= 20:
        warnings.append(
            {
                "level": "medium",
                "title": "部分项目进度滞后",
                "detail": f"当前不完整项目占比 {incomplete_rate}%。",
                "action": "按周跟踪缺失模块，要求提交可验证的改动证据。",
            }
        )

    if low_process_count > 0:
        warnings.append(
            {
                "level": "medium" if low_process_count <= 2 else "high",
                "title": "过程性学习证据不足",
                "detail": f"{low_process_count} 个项目在人机交互与迭代过程分偏低。",
                "action": "要求补充对话纪要、迭代记录和关键决策依据。",
            }
        )

    if not warnings:
        warnings.append(
            {
                "level": "low",
                "title": "当前无显著预警",
                "detail": "班级整体进度与过程表现稳定。",
                "action": "保持当前节奏，继续追踪关键节点。",
            }
        )

    warning_count = sum(1 for item in warnings if str(item.get("level") or "").lower() in {"high", "medium"})
    return {"warning_count": warning_count, "items": warnings}


def _safe_int(raw_value, default: int, minimum: int = 1, maximum: int = 200) -> int:
    try:
        value = int(raw_value)
    except Exception:
        value = default
    return max(minimum, min(value, maximum))


def _read_paging(request, default_page_size: int = 10, max_page_size: int = 50):
    params = request.query_params
    enabled = "page" in params or "page_size" in params
    page = _safe_int(params.get("page"), default=1, minimum=1, maximum=100000)
    page_size = _safe_int(params.get("page_size"), default=default_page_size, minimum=1, maximum=max_page_size)
    return enabled, page, page_size


def _paginate_queryset(queryset, page: int, page_size: int):
    total = queryset.count()
    start = (page - 1) * page_size
    return queryset[start: start + page_size], total


def _paginate_list(items: list, page: int, page_size: int):
    total = len(items)
    start = (page - 1) * page_size
    return items[start: start + page_size], total


def _teacher_student_ids(teacher: User, include_pending: bool = True, allow_plan_fallback: bool = False) -> list[int]:
    relation_ids = set(MentorshipRelation.objects.filter(teacher=teacher).values_list("student_id", flat=True))
    approved_ids = set(
        MentorshipApplication.objects.filter(
            teacher=teacher,
            status=MentorshipApplication.STATUS_APPROVED,
        ).values_list("student_id", flat=True)
    )

    # Auto-repair missing relation rows when approved applications exist.
    missing_relation_ids = approved_ids - relation_ids
    if missing_relation_ids:
        for student_id in missing_relation_ids:
            MentorshipRelation.objects.get_or_create(student_id=student_id, teacher=teacher)
        relation_ids = set(MentorshipRelation.objects.filter(teacher=teacher).values_list("student_id", flat=True))

    student_ids = set(relation_ids) | set(approved_ids)
    if include_pending:
        pending_ids = set(
            MentorshipApplication.objects.filter(
                teacher=teacher,
                status=MentorshipApplication.STATUS_PENDING,
            ).values_list("student_id", flat=True)
        )
        student_ids |= pending_ids

    if not student_ids and allow_plan_fallback:
        student_ids = set(
            Plan.objects.filter(status=Plan.STATUS_SUBMITTED)
            .values_list("student_id", flat=True)
            .distinct()[:400]
        )
    return sorted(student_ids)


def _teacher_can_access_plan(teacher: User, plan: Plan, allow_plan_fallback: bool = False) -> bool:
    if not plan:
        return False
    return int(plan.student_id) in set(
        _teacher_student_ids(
            teacher,
            include_pending=True,
            allow_plan_fallback=allow_plan_fallback,
        )
    )


def _normalize_project_type(project_type: str, plan_text: str = "") -> str:
    value = str(project_type or "").strip().lower()
    if value in {"public_welfare", "public", "nonprofit", "公益", "公益项目"}:
        return "public_welfare"
    if value in {"commercial", "business", "商业", "商业项目"}:
        return "commercial"
    text = str(plan_text or "").lower()
    welfare_keywords = ("公益", "志愿", "社会价值", "弱势", "公共服务", "非营利", "社会问题")
    return "public_welfare" if any(kw in text for kw in welfare_keywords) else "commercial"


def _parse_batch_users(raw_text: str, default_role: str = User.ROLE_STUDENT) -> list[dict]:
    rows = []
    safe_default_role = default_role if default_role in {User.ROLE_STUDENT, User.ROLE_TEACHER} else User.ROLE_STUDENT
    lines = [line.strip() for line in str(raw_text or "").splitlines() if line.strip()]
    for line in lines:
        if line.startswith("#"):
            continue

        # preferred format: 角色,姓名,学号/工号,默认密码,邮箱
        parts = [part.strip() for part in line.split(",")]
        if len(parts) < 2:
            continue

        role = safe_default_role
        full_name = ""
        account_no = ""
        password = ""
        email = ""

        if len(parts) >= 5:
            role = str(parts[0] or safe_default_role).strip().lower()
            full_name = parts[1]
            account_no = parts[2]
            password = parts[3]
            email = parts[4]
        else:
            # backward compatible fallback: username,password,role,email
            account_no = parts[0]
            password = parts[1] if len(parts) >= 2 else ""
            role = str(parts[2] or safe_default_role).strip().lower() if len(parts) >= 3 else safe_default_role
            email = parts[3] if len(parts) >= 4 else ""

        if role not in {User.ROLE_STUDENT, User.ROLE_TEACHER}:
            role = safe_default_role
        if not account_no or not password:
            continue

        rows.append(
            {
                "username": account_no,
                "account_no": account_no,
                "password": password,
                "role": role,
                "full_name": full_name,
                "email": email,
            }
        )
    return rows


def _history_to_context_text(history: list | None, *, max_items: int = 16, max_chars: int = 3200) -> str:
    rows = history if isinstance(history, list) else []
    chunks = []
    for item in rows[-max_items:]:
        if not isinstance(item, dict):
            continue
        role = str(item.get("role") or "").strip()
        content = str(item.get("content") or item.get("text") or "").strip()
        if not content:
            continue
        prefix = f"{role}: " if role else ""
        chunks.append(f"{prefix}{content}")
    return "\n".join(chunks)[:max_chars]


def _markdown_to_docx_bytes(title: str, markdown_text: str) -> bytes:
    if Document is None:
        return str(markdown_text or "").encode("utf-8", errors="ignore")
    doc = Document()
    doc.add_heading(title or "计划书导出", level=1)
    for raw_line in str(markdown_text or "").splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class AdminGraphObservabilityView(APIView):
    permission_classes = [IsPlatformAdmin]

    def get(self, request):
        excluded_ops = ["admin_hypergraph_preview", "admin_knowledge_graph_preview"]
        limit = int(request.query_params.get("limit") or 80)
        limit = max(20, min(limit, 200))
        page_enabled, page, page_size = _read_paging(request, default_page_size=20, max_page_size=100)

        logs_qs = GraphInvocationLog.objects.select_related("user", "plan").exclude(operation__in=excluded_ops).order_by("-created_at")

        keyword = str(request.query_params.get("q") or "").strip()
        graph_type = str(request.query_params.get("graph_type") or "").strip()
        operation = str(request.query_params.get("operation") or "").strip()
        source = str(request.query_params.get("source") or "").strip()
        success = str(request.query_params.get("success") or "").strip().lower()

        if keyword:
            logs_qs = logs_qs.filter(
                Q(operation__icontains=keyword)
                | Q(source__icontains=keyword)
                | Q(user__username__icontains=keyword)
                | Q(plan__title__icontains=keyword)
            )
        if graph_type:
            logs_qs = logs_qs.filter(graph_type=graph_type)
        if operation:
            logs_qs = logs_qs.filter(operation=operation)
        if source:
            logs_qs = logs_qs.filter(source=source)
        if success in {"true", "false"}:
            logs_qs = logs_qs.filter(success=(success == "true"))

        total_logs = logs_qs.count()
        failed_logs = logs_qs.filter(success=False).count()

        if page_enabled:
            logs_qs, _ = _paginate_queryset(logs_qs, page=page, page_size=page_size)
        else:
            logs_qs = logs_qs[:limit]
        logs = [
            {
                "id": row.id,
                "time": row.created_at,
                "graph_type": row.graph_type,
                "operation": row.operation,
                "source": row.source,
                "success": row.success,
                "user": row.user.username if row.user else "",
                "plan_id": row.plan_id,
                "plan_title": row.plan.title if row.plan else "",
                "detail": row.detail or {},
            }
            for row in logs_qs
        ]

        last_7_days = timezone.now() - timedelta(days=7)
        call_qs = GraphInvocationLog.objects.filter(created_at__gte=last_7_days).exclude(operation__in=excluded_ops)
        total_calls_7d = call_qs.count()
        success_calls_7d = call_qs.filter(success=True).count()
        hg_calls_7d = call_qs.filter(graph_type="hypergraph").count()
        kg_calls_7d = call_qs.filter(graph_type="knowledge_graph").count()

        review_qs = ReviewRecord.objects.order_by("-created_at")[:300]
        sample_size = review_qs.count()
        enabled_count = 0
        with_provenance_count = 0
        with_hypergraph_suggestion_count = 0
        explainable_review_count = 0
        explainable_item_numerator = 0
        explainable_item_denominator = 0

        for review in review_qs:
            meta = (review.review_meta or {})
            hypergraph = (meta.get("hypergraph") or {})
            provenance = (hypergraph.get("provenance") or {})
            node_sources = provenance.get("node_sources") or []
            edge_evidence = provenance.get("hyperedge_evidence") or []
            suggestions = review.suggestions or []

            if hypergraph.get("enabled"):
                enabled_count += 1
            if node_sources or edge_evidence:
                with_provenance_count += 1
            if any(("[瓒呭浘鏍￠獙]" in str(s or "") or "[瓒呭浘琛ラ綈]" in str(s or "")) for s in suggestions):
                with_hypergraph_suggestion_count += 1

            anchors = []
            for item in node_sources:
                anchors.extend([
                    str(item.get("text") or "").strip(),
                    str(item.get("node_name") or "").strip(),
                ])
            for item in edge_evidence:
                anchors.append(str(item.get("text") or "").strip())
            anchors = [item for item in anchors if item]

            targets = []
            for issue in (review.issues or []):
                targets.extend([
                    _safe_review_item_text(issue, "description", "title", "text"),
                    _safe_review_item_text(issue, "snippet", "anchor", "quote"),
                ])
            for ann in (review.annotations or []):
                targets.extend([
                    _safe_review_item_text(ann, "description", "title", "text"),
                    _safe_review_item_text(ann, "question", "guiding_question"),
                ])
            targets = [item for item in targets if item]

            if not targets:
                continue

            matched_targets = 0
            for target in targets:
                if any(_text_hit(anchor, target) for anchor in anchors):
                    matched_targets += 1

            explainable_item_numerator += matched_targets
            explainable_item_denominator += len(targets)
            if matched_targets > 0:
                explainable_review_count += 1

        return ok(
            {
                "recent_logs": logs,
                "log_stats": {
                    "total": total_logs,
                    "failed": failed_logs,
                    "failure_rate": _safe_ratio(failed_logs, total_logs),
                },
                "pagination": {
                    "enabled": page_enabled,
                    "page": page,
                    "page_size": page_size,
                    "total": total_logs,
                },
                "call_stats": {
                    "window_days": 7,
                    "total_calls": total_calls_7d,
                    "success_calls": success_calls_7d,
                    "success_rate": _safe_ratio(success_calls_7d, total_calls_7d),
                    "hypergraph_calls": hg_calls_7d,
                    "knowledge_graph_calls": kg_calls_7d,
                },
                "feedback_stats": {
                    "sample_size": sample_size,
                    "enabled_count": enabled_count,
                    "enabled_rate": _safe_ratio(enabled_count, sample_size),
                    "with_provenance_count": with_provenance_count,
                    "provenance_rate": _safe_ratio(with_provenance_count, sample_size),
                    "with_hypergraph_suggestion_count": with_hypergraph_suggestion_count,
                    "suggestion_hit_rate": _safe_ratio(with_hypergraph_suggestion_count, sample_size),
                    "explainable_review_count": explainable_review_count,
                    "review_explainability_rate": _safe_ratio(explainable_review_count, sample_size),
                    "explainable_item_rate": _safe_ratio(explainable_item_numerator, explainable_item_denominator),
                },
            }
        )


class TeacherListView(generics.ListAPIView):
    serializer_class = UserSerializer

    def get_queryset(self):
        keyword = str(self.request.query_params.get("q") or "").strip()
        queryset = User.objects.filter(
            role=User.ROLE_TEACHER,
            is_active=True,
            is_staff=False,
            is_superuser=False,
        )
        if keyword:
            queryset = queryset.filter(
                Q(username__icontains=keyword)
                | Q(full_name__icontains=keyword)
                | Q(organization__icontains=keyword)
                | Q(coaching_direction__icontains=keyword)
                | Q(subject__icontains=keyword)
            )
        return queryset.order_by("username")

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        page_enabled, page, page_size = _read_paging(request, default_page_size=8, max_page_size=50)
        if page_enabled:
            page_queryset, total = _paginate_queryset(queryset, page=page, page_size=page_size)
            data = UserSerializer(page_queryset, many=True).data
            return ok({"items": data, "total": total, "page": page, "page_size": page_size})
        data = UserSerializer(queryset, many=True).data
        return ok(data)


class ChatGuideView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        stage = request.data.get("stage", "idea")
        answer = request.data.get("answer", "")
        session_id = request.data.get("session_id") or "default"
        result = next_guiding_question(stage, answer)

        record = ConversationRecord.objects.create(
            user=request.user,
            stage=result["stage"],
            question=result["question"],
            answer=answer,
            metadata={
                "question_type": result.get("question_type", "guided"),
                "scene_prompt": result.get("scene_prompt", ""),
                "chat_title": result.get("chat_title", ""),
                "session_id": str(session_id),
            },
        )
        log_event("chat_guide", {"user": request.user.username, "stage": result["stage"]})
        return ok({"next": result, "record": ConversationRecordSerializer(record).data})


class LogicValidateView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        text = request.data.get("text", "")
        if not text.strip():
            return ok(
                {
                    "issues": [
                        {
                            "position": "input",
                            "description": "请输入要校验的内容",
                            "question": "你最想先验证哪个部分？",
                        }
                    ]
                }
            )
        return ok(validate_logic(text))


class PlanUploadView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        if not request.FILES.get("pdf_file"):
            return fail("请先选择 PDF 或 Word 文件")

        uploaded = request.FILES.get("pdf_file")
        ext = Path(uploaded.name).suffix.lower()
        if ext not in {".pdf", ".doc", ".docx"}:
            return fail("仅支持 PDF、DOC、DOCX 文件")
        if uploaded.size > 50 * 1024 * 1024:
            return fail("文件大小不能超过 50MB")

        base_plan_id = request.data.get("base_plan_id")
        base_plan = None
        if str(base_plan_id or "").isdigit():
            base_plan = Plan.objects.filter(id=int(base_plan_id), student=request.user).first()

        title = (base_plan.title if base_plan else request.data.get("title", "")).strip()
        if not title:
            return fail("方案名称不能为空")
        latest = Plan.objects.filter(student=request.user, title=title).order_by("-version").first()
        version = latest.version + 1 if latest else 1

        plan = None
        try:
            plan = Plan.objects.create(
                student=request.user,
                title=title,
                version=version,
                file_size=uploaded.size,
                pdf_file=uploaded,
            )

            if ext in {".doc", ".docx"}:
                # Parse Word text first for review stability, then convert to PDF for page-style preview.
                parsed_word = _parse_plan_file(plan.pdf_file.path, ext)
                original_word_path = plan.pdf_file.path
                converted_ok = False
                with tempfile.TemporaryDirectory() as temp_dir:
                    converted_pdf = _run_office_to_pdf(plan.pdf_file.path, temp_dir)
                    if converted_pdf:
                        with open(converted_pdf, "rb") as fp:
                            output_name = f"{Path(uploaded.name).stem}_v{version}.pdf"
                            plan.pdf_file.save(output_name, File(fp), save=False)
                        converted_ok = True

                if converted_ok:
                    if original_word_path and os.path.exists(original_word_path):
                        try:
                            os.remove(original_word_path)
                        except Exception:
                            pass
                    ext = ".pdf"
                    parsed_pdf = _parse_plan_file(plan.pdf_file.path, ext)
                    extracted_text = (parsed_word.get("text") or "").strip() or (parsed_pdf.get("text") or "").strip()
                    pages = parsed_pdf.get("pages") or _build_pages_from_text(extracted_text)
                    parsed = {
                        "text": extracted_text,
                        "pages": pages,
                        "page_count": int(parsed_pdf.get("page_count") or len(pages) or 0),
                    }
                else:
                    # Converter unavailable: keep original Word file and rely on extracted text for review.
                    parsed = {
                        "text": (parsed_word.get("text") or "").strip(),
                        "pages": parsed_word.get("pages") or _build_pages_from_text((parsed_word.get("text") or "").strip()),
                        "page_count": int(parsed_word.get("page_count") or 0),
                    }
            else:
                parsed = _parse_plan_file(plan.pdf_file.path, ext)

            if ext == ".pdf" and parsed["page_count"] > 100:
                plan.pdf_file.delete(save=False)
                plan.delete()
                return fail("PDF页数不能超过100页")

            if parsed["page_count"] == 0 or not (parsed.get("text") or "").strip():
                plan.pdf_file.delete(save=False)
                plan.delete()
                return fail("鏂囨。瑙ｆ瀽澶辫触锛岃涓婁紶鍙В鏋愭枃鏈殑PDF鎴朩ord")

            plan.extracted_text = parsed["text"]
            plan.page_count = parsed["page_count"]
            plan.parsed_pages = parsed["pages"]
            plan.save(update_fields=["extracted_text", "page_count", "file_size", "parsed_pages"])
        except Exception:
            if plan:
                plan.pdf_file.delete(save=False)
                plan.delete()
            return fail("文档解析失败，请检查文件是否损坏")

        log_event("plan_upload", {"user": request.user.username, "plan_id": plan.id, "version": version})
        return ok(PlanSerializer(plan).data, message="涓婁紶鎴愬姛", status_code=status.HTTP_201_CREATED)


class PlanListView(generics.ListAPIView):
    serializer_class = PlanSerializer

    def get_queryset(self):
        user = self.request.user
        if user.role == User.ROLE_TEACHER:
            student_ids = _teacher_student_ids(user, include_pending=True, allow_plan_fallback=True)
            return Plan.objects.filter(student_id__in=student_ids, status=Plan.STATUS_SUBMITTED)
        return Plan.objects.filter(student=user)

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset().order_by("-created_at")
        keyword = str(request.query_params.get("q") or "").strip()
        status_filter = str(request.query_params.get("status") or "").strip().lower()

        if keyword:
            queryset = queryset.filter(title__icontains=keyword)
        if status_filter in {Plan.STATUS_DRAFT, Plan.STATUS_SUBMITTED}:
            queryset = queryset.filter(status=status_filter)

        page_enabled, page, page_size = _read_paging(request, default_page_size=8, max_page_size=50)
        if page_enabled:
            page_queryset, total = _paginate_queryset(queryset, page=page, page_size=page_size)
            data = PlanSerializer(page_queryset, many=True).data
            return ok({"items": data, "total": total, "page": page, "page_size": page_size})

        data = PlanSerializer(queryset, many=True).data
        return ok(data)


class PlanDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, plan_id):
        plan = get_object_or_404(Plan, id=plan_id)
        scope = str(request.query_params.get("scope") or "single").strip().lower()

        if request.user.role == User.ROLE_STUDENT:
            if plan.student_id != request.user.id:
                return fail("鏃犳潈闄愬垹闄よ鏂规", status.HTTP_403_FORBIDDEN)
            if scope in {"project", "all", "all_versions"}:
                same_project_qs = Plan.objects.filter(student=request.user, title=plan.title)
                deleted_count = same_project_qs.count()
                for item in list(same_project_qs):
                    _delete_plan_with_local_file(item)
                return ok({"deleted_count": deleted_count, "scope": "project"}, message="项目全部版本已删除")
        elif request.user.role == User.ROLE_TEACHER:
            if not _teacher_can_access_plan(request.user, plan, allow_plan_fallback=True):
                return fail("鏃犳潈闄愬垹闄よ鏂规", status.HTTP_403_FORBIDDEN)
        else:
            return fail("鏃犳潈闄愬垹闄よ鏂规", status.HTTP_403_FORBIDDEN)

        _delete_plan_with_local_file(plan)
        return ok({"deleted_count": 1, "scope": "single"}, message="删除成功")


class PlanSubmitToTeacherView(APIView):
    permission_classes = [IsStudent]

    def post(self, request, plan_id):
        plan = get_object_or_404(Plan, id=plan_id, student=request.user)
        if plan.status == Plan.STATUS_SUBMITTED:
            return ok({"plan": PlanSerializer(plan).data}, message="该方案已提交给教师")
        teacher_id = request.data.get("teacher_id")
        note = request.data.get("note", "")
        if not str(teacher_id or "").isdigit():
            return fail("teacher_id 参数非法")
        teacher = User.objects.filter(id=int(teacher_id), role=User.ROLE_TEACHER).first()
        if not teacher:
            return fail("未找到目标教师")

        relation = MentorshipRelation.objects.filter(student=request.user, teacher=teacher).first()
        if not relation:
            approved_exists = MentorshipApplication.objects.filter(
                student=request.user,
                teacher=teacher,
                status=MentorshipApplication.STATUS_APPROVED,
            ).exists()
            if not approved_exists:
                return fail("你与该教师尚未建立师徒关系")
            relation, _ = MentorshipRelation.objects.get_or_create(student=request.user, teacher=teacher)

        plan.status = Plan.STATUS_SUBMITTED
        plan.note = note
        plan.save(update_fields=["status", "note"])

        MessageNotification.objects.create(
            user=relation.teacher,
            title="收到新方案",
            content=f"学生{request.user.display_name}提交了方案《{plan.title}》V{plan.version}。",
        )
        return ok({"plan": PlanSerializer(plan).data, "teacher_id": teacher.id}, message="提交成功")


class ReviewGenerateView(APIView):
    def post(self, request):
        plan_id = request.data.get("plan_id")
        audience_role = request.data.get("audience_role", request.user.role)
        plan = get_object_or_404(Plan, id=plan_id)

        if request.user.role == User.ROLE_STUDENT and plan.student_id != request.user.id:
            return fail("鏃犳潈闄愯闂鏂规", status.HTTP_403_FORBIDDEN)

        if audience_role not in {"student", "teacher"}:
            return fail("audience_role鍙傛暟闈炴硶")

        if request.user.role == User.ROLE_STUDENT:
            audience_role = "student"
        elif request.user.role == User.ROLE_TEACHER:
            if not _teacher_can_access_plan(request.user, plan, allow_plan_fallback=True):
                return fail("鏃犳潈闄愯闂鏂规", status.HTTP_403_FORBIDDEN)
            audience_role = "teacher"

        pages = _ensure_plan_pages(plan)

        if not pages:
            pages = [{"page": 1, "text": plan.extracted_text or plan.title}]

        full_text = "\n".join((item or {}).get("text", "") for item in pages).strip()
        if not full_text:
            return fail("无法解析文档文本，请重新上传可解析版本")

        project_type = _normalize_project_type(request.data.get("project_type"), plan_text=full_text)
        hg_report = build_plan_diagnosis(plan)
        payload = diagnose_pdf(
            pages,
            audience_role,
            hypergraph_context=hg_report,
            project_type=project_type,
        )
        payload.setdefault("review_meta", {})
        payload["review_meta"]["parsed_text_length"] = len(full_text)
        payload["review_meta"]["parsed_page_count"] = len(pages)
        payload["review_meta"]["project_type"] = project_type
        payload["review_meta"]["hypergraph"] = _build_hypergraph_meta(hg_report)
        if hg_report.get("enabled"):
            hg_suggestions = _build_hypergraph_suggestions(hg_report)
            if hg_suggestions:
                payload["suggestions"] = (payload.get("suggestions") or []) + hg_suggestions

        review = ReviewRecord.objects.create(
            plan=plan,
            reviewer=request.user,
            audience_role=audience_role,
            issues=payload["issues"],
            annotations=payload.get("annotations", []),
            guidance_questions=payload["guidance_questions"],
            examples=payload.get("examples", []),
            suggestions=payload["suggestions"],
            review_meta=payload.get("review_meta", {}),
            summary=payload["summary"],
        )
        return ok(ReviewRecordSerializer(review).data, message="鎵归槄瀹屾垚")


def _run_review_task(task_id: int, reviewer_id: int, plan_id: int, audience_role: str, project_type: str = "auto"):
    close_old_connections()
    try:
        task = ConversationRecord.objects.get(id=task_id)
        plan = Plan.objects.get(id=plan_id)
        reviewer = User.objects.get(id=reviewer_id)

        task.metadata = {
            **(task.metadata or {}),
            "status": "running",
            "started_at": timezone.now().isoformat(),
        }
        task.save(update_fields=["metadata"])

        pages = _ensure_plan_pages(plan)

        if not pages:
            pages = [{"page": 1, "text": plan.extracted_text or plan.title}]

        full_text = "\n".join((item or {}).get("text", "") for item in pages).strip()
        if not full_text:
            raise ValueError("无法解析文档文本，请重新上传可解析版本")

        normalized_project_type = _normalize_project_type(project_type, plan_text=full_text)
        hg_report = build_plan_diagnosis(plan)
        payload = diagnose_pdf(
            pages,
            audience_role,
            hypergraph_context=hg_report,
            project_type=normalized_project_type,
        )
        payload.setdefault("review_meta", {})
        payload["review_meta"]["parsed_text_length"] = len(full_text)
        payload["review_meta"]["parsed_page_count"] = len(pages)
        payload["review_meta"]["project_type"] = normalized_project_type
        payload["review_meta"]["hypergraph"] = _build_hypergraph_meta(hg_report)
        if hg_report.get("enabled"):
            hg_suggestions = _build_hypergraph_suggestions(hg_report)
            if hg_suggestions:
                payload["suggestions"] = (payload.get("suggestions") or []) + hg_suggestions

        review = ReviewRecord.objects.create(
            plan=plan,
            reviewer=reviewer,
            audience_role=audience_role,
            issues=payload["issues"],
            annotations=payload.get("annotations", []),
            guidance_questions=payload["guidance_questions"],
            examples=payload.get("examples", []),
            suggestions=payload["suggestions"],
            review_meta=payload.get("review_meta", {}),
            summary=payload["summary"],
        )

        task.answer = payload.get("summary", "")
        task.metadata = {
            **(task.metadata or {}),
            "status": "done",
            "review_id": review.id,
        }
        task.save(update_fields=["answer", "metadata"])
    except Exception as exc:
        task = ConversationRecord.objects.filter(id=task_id).first()
        if task:
            task.metadata = {
                **(task.metadata or {}),
                "status": "error",
                "error": str(exc),
            }
            task.save(update_fields=["metadata"])
    finally:
        close_old_connections()


class ReviewGenerateAsyncView(APIView):
    def post(self, request):
        plan_id = request.data.get("plan_id")
        audience_role = request.data.get("audience_role", request.user.role)
        plan = get_object_or_404(Plan, id=plan_id)

        if request.user.role == User.ROLE_STUDENT and plan.student_id != request.user.id:
            return fail("鏃犳潈闄愯闂鏂规", status.HTTP_403_FORBIDDEN)

        if audience_role not in {"student", "teacher"}:
            return fail("audience_role鍙傛暟闈炴硶")

        if request.user.role == User.ROLE_STUDENT:
            audience_role = "student"
        elif request.user.role == User.ROLE_TEACHER:
            if not _teacher_can_access_plan(request.user, plan, allow_plan_fallback=True):
                return fail("鏃犳潈闄愯闂鏂规", status.HTTP_403_FORBIDDEN)
            audience_role = "teacher"

        project_type = _normalize_project_type(request.data.get("project_type"), plan_text=plan.extracted_text or "")

        task = ConversationRecord.objects.create(
            user=request.user,
            stage="review_task",
            question=f"review_plan_{plan.id}",
            answer="",
            metadata={
                "status": "queued",
                "plan_id": plan.id,
                "audience_role": audience_role,
                "project_type": project_type,
                "queued_at": timezone.now().isoformat(),
            },
        )

        thread = threading.Thread(
            target=_run_review_task,
            args=(task.id, request.user.id, plan.id, audience_role, project_type),
            daemon=True,
        )
        thread.start()
        return ok({"task_id": task.id, "status": "queued"}, message="审阅任务已提交")


class ReviewTaskStatusView(APIView):
    def get(self, request, task_id):
        task = get_object_or_404(ConversationRecord, id=task_id, stage="review_task")
        plan_id = (task.metadata or {}).get("plan_id")
        plan = get_object_or_404(Plan, id=plan_id)

        if request.user.role == User.ROLE_STUDENT and plan.student_id != request.user.id:
            return fail("鏃犳潈闄愯闂浠诲姟", status.HTTP_403_FORBIDDEN)
        if request.user.role == User.ROLE_TEACHER:
            if not _teacher_can_access_plan(request.user, plan, allow_plan_fallback=True):
                return fail("鏃犳潈闄愯闂浠诲姟", status.HTTP_403_FORBIDDEN)

        task_meta = (task.metadata or {}).copy()
        task_status = str(task_meta.get("status") or "queued").lower()
        if task_status in {"queued", "running"}:
            started_raw = task_meta.get("started_at") or task_meta.get("queued_at")
            started_at = parse_datetime(str(started_raw)) if started_raw else None
            if started_at and timezone.is_naive(started_at):
                started_at = timezone.make_aware(started_at, timezone.get_current_timezone())
            elapsed = (timezone.now() - (started_at or task.created_at)).total_seconds()
            if elapsed > REVIEW_TASK_TIMEOUT_SECONDS:
                task_status = "error"
                task_meta["status"] = task_status
                task_meta["error"] = (
                    f"审阅任务超时（>{REVIEW_TASK_TIMEOUT_SECONDS}秒）。"
                    "请稍后重试，或先缩短方案内容后再发起审阅。"
                )
                task.metadata = task_meta
                task.save(update_fields=["metadata"])

        review_data = None
        if task_status == "done" and task_meta.get("review_id"):
            review = ReviewRecord.objects.filter(id=task_meta.get("review_id"), plan=plan).first()
            if review:
                review_data = ReviewRecordSerializer(review).data

        return ok(
            {
                "task_id": task.id,
                "status": task_status,
                "error": task_meta.get("error", ""),
                "review": review_data,
            }
        )


class StudentPdfChatView(APIView):
    permission_classes = [IsStudent]

    def get(self, request):
        plan_id = request.query_params.get("plan_id")
        if not plan_id:
            return fail("plan_id涓嶈兘涓虹┖")
        plan = get_object_or_404(Plan, id=plan_id, student=request.user)
        records = ConversationRecord.objects.filter(
            user=request.user,
            stage="pdf_chat",
            metadata__plan_id=plan.id,
        ).order_by("created_at")
        ordered = []
        for item in records:
            if item.question:
                ordered.append({"role": "student", "content": item.question})
            if item.answer:
                ordered.append({"role": "ai", "content": item.answer})
        return ok({"messages": ordered})

    def post(self, request):
        plan_id = request.data.get("plan_id")
        question = (request.data.get("question") or "").strip()
        if not plan_id:
            return fail("plan_id涓嶈兘涓虹┖")
        if not question:
            return fail("闂涓嶈兘涓虹┖")

        plan = get_object_or_404(Plan, id=plan_id, student=request.user)
        if not (plan.extracted_text or "").strip():
            try:
                _ensure_plan_pages(plan)
            except Exception:
                pass

        review_id = request.data.get("review_id")
        latest_review = None
        if review_id:
            latest_review = ReviewRecord.objects.filter(
                id=review_id,
                plan=plan,
                audience_role=ReviewRecord.ROLE_STUDENT,
            ).first()
        if not latest_review:
            latest_review = ReviewRecord.objects.filter(plan=plan, audience_role=ReviewRecord.ROLE_STUDENT).first()

        recent_records = list(
            ConversationRecord.objects.filter(
                user=request.user,
                stage="pdf_chat",
                metadata__plan_id=plan.id,
            )[:6]
        )
        context = {
            "plan_id": plan.id,
            "title": plan.title,
            "page_count": plan.page_count,
            "plan_excerpt": (plan.extracted_text or "")[:2200],
            "review_summary": latest_review.summary if latest_review else "",
            "review_issues": [
                _safe_review_item_text(item, "description", "title", "text")
                for item in ((latest_review.issues if latest_review else []) or [])
                if _safe_review_item_text(item, "description", "title", "text")
            ],
            "review_questions": (latest_review.guidance_questions if latest_review else []) or [],
            "hypergraph": ((latest_review.review_meta if latest_review else {}) or {}).get("hypergraph") or build_plan_diagnosis(plan),
            "chat_history": [
                {"question": record.question, "answer": record.answer}
                for record in recent_records
            ],
        }
        answer = student_pdf_chat_answer(question, context)
        ConversationRecord.objects.create(
            user=request.user,
            stage="pdf_chat",
            question=question,
            answer=answer,
            metadata={"context_type": "student_pdf_chat", "plan_id": plan.id, "review_id": latest_review.id if latest_review else None},
        )
        return ok({"answer": answer, "context": context})


class PlanReviewListView(generics.ListAPIView):
    serializer_class = ReviewRecordSerializer

    def get_queryset(self):
        plan_id = self.kwargs["plan_id"]
        plan = get_object_or_404(Plan, id=plan_id)
        user = self.request.user
        if user.role == User.ROLE_STUDENT and plan.student_id != user.id:
            return ReviewRecord.objects.none()
        if user.role == User.ROLE_TEACHER:
            if not _teacher_can_access_plan(user, plan, allow_plan_fallback=True):
                return ReviewRecord.objects.none()
        return ReviewRecord.objects.filter(plan=plan)

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset().order_by("-created_at")
        keyword = str(request.query_params.get("q") or "").strip()
        status_filter = str(request.query_params.get("status") or "").strip().lower()
        if keyword:
            queryset = queryset.filter(Q(summary__icontains=keyword) | Q(plan__title__icontains=keyword))
        if status_filter in {ReviewRecord.ROLE_STUDENT, ReviewRecord.ROLE_TEACHER}:
            queryset = queryset.filter(audience_role=status_filter)

        page_enabled, page, page_size = _read_paging(request, default_page_size=8, max_page_size=50)
        if page_enabled:
            page_queryset, total = _paginate_queryset(queryset, page=page, page_size=page_size)
            data = ReviewRecordSerializer(page_queryset, many=True).data
            return ok({"items": data, "total": total, "page": page, "page_size": page_size})
        data = ReviewRecordSerializer(queryset, many=True).data
        return ok(data)


class MentorshipApplyView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        payload = request.data.copy()
        if not payload.get("teacher") and payload.get("teacher_id"):
            payload["teacher"] = payload.get("teacher_id")

        serializer = MentorshipApplicationSerializer(data=payload)
        serializer.is_valid(raise_exception=True)
        teacher = serializer.validated_data["teacher"]
        if teacher.role != User.ROLE_TEACHER:
            return fail("鍙兘鐢宠鏁欏笀鐢ㄦ埛")

        application, created = MentorshipApplication.objects.update_or_create(
            student=request.user,
            teacher=teacher,
            defaults={
                "startup_direction": serializer.validated_data["startup_direction"],
                "demand_note": serializer.validated_data["demand_note"],
                "status": MentorshipApplication.STATUS_PENDING,
                "reject_reason": "",
            },
        )

        MessageNotification.objects.create(
            user=teacher,
            title="鏂扮殑鎸囧鐢宠",
            content=f"学生{request.user.username}申请指导：{application.startup_direction}",
        )
        code = status.HTTP_201_CREATED if created else status.HTTP_200_OK
        return ok(MentorshipApplicationSerializer(application).data, status_code=code)


class MentorshipApplicationListView(generics.ListAPIView):
    serializer_class = MentorshipApplicationSerializer

    def get_queryset(self):
        user = self.request.user
        if user.role == User.ROLE_TEACHER:
            return MentorshipApplication.objects.filter(teacher=user)
        return MentorshipApplication.objects.filter(student=user)

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset().select_related("student", "teacher").order_by("-created_at")
        keyword = str(request.query_params.get("q") or "").strip()
        status_filter = str(request.query_params.get("status") or "").strip().lower()
        if keyword:
            queryset = queryset.filter(
                Q(student__username__icontains=keyword)
                | Q(teacher__username__icontains=keyword)
                | Q(startup_direction__icontains=keyword)
            )
        if status_filter in {
            MentorshipApplication.STATUS_PENDING,
            MentorshipApplication.STATUS_APPROVED,
            MentorshipApplication.STATUS_REJECTED,
        }:
            queryset = queryset.filter(status=status_filter)

        page_enabled, page, page_size = _read_paging(request, default_page_size=8, max_page_size=50)
        if page_enabled:
            page_queryset, total = _paginate_queryset(queryset, page=page, page_size=page_size)
            data = MentorshipApplicationSerializer(page_queryset, many=True).data
            return ok({"items": data, "total": total, "page": page, "page_size": page_size})
        data = MentorshipApplicationSerializer(queryset, many=True).data
        return ok(data)


class MentorshipAuditView(APIView):
    permission_classes = [IsTeacher]

    def post(self, request, application_id):
        action = request.data.get("action", "")
        reject_reason = request.data.get("reject_reason", "")
        app = get_object_or_404(MentorshipApplication, id=application_id, teacher=request.user)
        if app.status != MentorshipApplication.STATUS_PENDING:
            return fail("该申请已处理，状态不可再修改")

        if action == "approve":
            app.status = MentorshipApplication.STATUS_APPROVED
            app.reject_reason = ""
            MentorshipRelation.objects.get_or_create(student=app.student, teacher=app.teacher)
            MessageNotification.objects.create(
                user=app.student,
                title="鎸囧鐢宠宸查€氳繃",
                content=f"教师{request.user.username}已通过你的指导申请。",
            )
        elif action == "reject":
            app.status = MentorshipApplication.STATUS_REJECTED
            app.reject_reason = reject_reason or "当前名额已满"
            MessageNotification.objects.create(
                user=app.student,
                title="指导申请未通过",
                content=f"教师{request.user.username}拒绝申请：{app.reject_reason}",
            )
        else:
            return fail("action蹇呴』涓篴pprove鎴杛eject")

        app.audit_action = action
        app.audited_at = timezone.now()
        app.audited_by = request.user
        app.save(update_fields=["status", "reject_reason", "audit_action", "audited_at", "audited_by"])
        return ok(MentorshipApplicationSerializer(app).data)


class TeacherStudentListView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        keyword = str(request.query_params.get("q") or "").strip()
        page_enabled, page, page_size = _read_paging(request, default_page_size=10, max_page_size=50)
        student_ids = _teacher_student_ids(request.user, include_pending=False, allow_plan_fallback=False)
        if not student_ids:
            if page_enabled:
                return ok({"items": [], "total": 0, "page": page, "page_size": page_size})
            return ok([])

        queryset = User.objects.filter(id__in=student_ids, role=User.ROLE_STUDENT).order_by("username")
        if keyword:
            queryset = queryset.filter(
                Q(username__icontains=keyword)
                | Q(full_name__icontains=keyword)
                | Q(student_no__icontains=keyword)
                | Q(major__icontains=keyword)
            )

        relation_map = {
            row.student_id: row
            for row in MentorshipRelation.objects.filter(teacher=request.user, student_id__in=student_ids).select_related("student")
        }
        app_rows = (
            MentorshipApplication.objects.filter(teacher=request.user, student_id__in=student_ids)
            .order_by("student_id", "-created_at")
        )
        latest_app_map = {}
        for row in app_rows:
            latest_app_map.setdefault(row.student_id, row)

        def _serialize(student: User):
            relation = relation_map.get(student.id)
            latest_app = latest_app_map.get(student.id)
            return {
                "student_id": student.id,
                "student_name": student.display_name,
                "username": student.username,
                "student_no": student.student_no or student.username,
                "major": student.major or "",
                "status": "bound" if relation else "unbound",
                "status_text": "已绑定" if relation else "未绑定",
                "bound_at": relation.created_at if relation else None,
                "application_status": latest_app.status if latest_app else "",
                "application_status_text": {
                    MentorshipApplication.STATUS_APPROVED: "已通过",
                    MentorshipApplication.STATUS_REJECTED: "已拒绝",
                    MentorshipApplication.STATUS_PENDING: "待处理",
                }.get(getattr(latest_app, "status", ""), ""),
            }

        if page_enabled:
            page_queryset, total = _paginate_queryset(queryset, page=page, page_size=page_size)
            data = [_serialize(student) for student in page_queryset]
            return ok({"items": data, "total": total, "page": page, "page_size": page_size})

        data = [_serialize(student) for student in queryset]
        return ok(data)


class TeacherStudentUnbindView(APIView):
    permission_classes = [IsTeacher]

    def post(self, request, student_id):
        student = get_object_or_404(User, id=student_id, role=User.ROLE_STUDENT)
        relation = MentorshipRelation.objects.filter(teacher=request.user, student=student).first()
        if not relation:
            return fail("该学生当前未与您绑定")

        with transaction.atomic():
            relation.delete()
            MentorshipApplication.objects.filter(
                student=student,
                teacher=request.user,
                status=MentorshipApplication.STATUS_APPROVED,
            ).update(
                status=MentorshipApplication.STATUS_REJECTED,
                reject_reason="教师解除绑定",
                audit_action="unbind",
                audited_at=timezone.now(),
                audited_by=request.user,
            )

        MessageNotification.objects.create(
            user=student,
            title="师生关系已解除",
            content=f"教师{request.user.display_name}已解除与你的指导关系，如需继续指导可重新发起申请。",
        )
        return ok({"student_id": student.id, "unbound": True}, message="解绑成功")


class MessageListView(generics.ListAPIView):
    serializer_class = MessageNotificationSerializer

    def get_queryset(self):
        return MessageNotification.objects.filter(user=self.request.user)

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset().order_by("-created_at")
        keyword = str(request.query_params.get("q") or "").strip()
        status_filter = str(request.query_params.get("status") or "").strip().lower()

        if keyword:
            queryset = queryset.filter(Q(title__icontains=keyword) | Q(content__icontains=keyword))
        if status_filter == "read":
            queryset = queryset.filter(is_read=True)
        elif status_filter == "unread":
            queryset = queryset.filter(is_read=False)

        page_enabled, page, page_size = _read_paging(request, default_page_size=8, max_page_size=50)
        if page_enabled:
            page_queryset, total = _paginate_queryset(queryset, page=page, page_size=page_size)
            data = self.get_serializer(page_queryset, many=True).data
            return ok(
                {
                    "items": data,
                    "total": total,
                    "page": page,
                    "page_size": page_size,
                    "timezone": "UTC+8",
                    "format": "YYYY-MM-DD HH:MM:SS",
                }
            )

        data = self.get_serializer(queryset, many=True).data
        return ok({"items": data, "timezone": "UTC+8", "format": "YYYY-MM-DD HH:MM:SS"})


class MessageReadView(APIView):
    def post(self, request):
        message_ids = request.data.get("message_ids") or []
        if isinstance(message_ids, int):
            message_ids = [message_ids]
        if not isinstance(message_ids, list):
            return fail("message_ids鍙傛暟鏍煎紡閿欒")

        queryset = MessageNotification.objects.filter(user=request.user, id__in=message_ids)
        updated = queryset.update(is_read=True)
        return ok({"updated": updated}, message="已更新已读状态")


class HistoryView(APIView):
    def get(self, request):
        plans_qs = Plan.objects.filter(student=request.user).order_by("-created_at")[:100]
        reviews_qs = ReviewRecord.objects.filter(plan__student=request.user).select_related("plan").order_by("-created_at")[:100]

        plans = []
        for item in plans_qs:
            plans.append(
                {
                    "id": item.id,
                    "title": item.title,
                    "version": item.version,
                    "status": item.status,
                    "has_ai_review": item.reviews.exists(),
                    "created_at": item.created_at,
                }
            )

        reviews = []
        for item in reviews_qs:
            reviews.append(
                {
                    "id": item.id,
                    "plan_id": item.plan_id,
                    "plan_title": item.plan.title,
                    "created_at": item.created_at,
                    "summary": item.summary,
                }
            )

        submissions = []
        for item in plans_qs:
            if item.status == Plan.STATUS_SUBMITTED:
                submissions.append(
                    {
                        "plan_id": item.id,
                        "title": item.title,
                        "version": item.version,
                        "created_at": item.created_at,
                    }
                )

        keyword = str(request.query_params.get("q") or "").strip().lower()
        type_filter = str(request.query_params.get("status") or request.query_params.get("type") or "").strip().lower()
        stream_mode = (str(request.query_params.get("mode") or "").strip().lower() == "stream")
        page_enabled, page, page_size = _read_paging(request, default_page_size=8, max_page_size=50)

        if stream_mode or page_enabled:
            stream_items = []
            for item in plans:
                stream_items.append(
                    {
                        "id": f"plan-{item.get('id')}",
                        "type": "plan",
                        "action": f"上传方案《{item.get('title') or '-'}》v{item.get('version') or 1}",
                        "title": item.get("title") or "",
                        "status": item.get("status") or "",
                        "created_at": item.get("created_at"),
                    }
                )
            for item in reviews:
                stream_items.append(
                    {
                        "id": f"review-{item.get('id')}",
                        "type": "review",
                        "action": f"生成审阅《{item.get('plan_title') or '-'}》",
                        "title": item.get("plan_title") or "",
                        "summary": item.get("summary") or "",
                        "created_at": item.get("created_at"),
                    }
                )
            for item in submissions:
                stream_items.append(
                    {
                        "id": f"submission-{item.get('plan_id')}-{item.get('version')}",
                        "type": "submission",
                        "action": f"提交方案《{item.get('title') or '-'}》v{item.get('version') or 1}",
                        "title": item.get("title") or "",
                        "status": Plan.STATUS_SUBMITTED,
                        "created_at": item.get("created_at"),
                    }
                )

            if type_filter in {"plan", "review", "submission"}:
                stream_items = [item for item in stream_items if item.get("type") == type_filter]
            if keyword:
                stream_items = [
                    item
                    for item in stream_items
                    if keyword in str(item.get("action") or "").lower()
                    or keyword in str(item.get("title") or "").lower()
                    or keyword in str(item.get("summary") or "").lower()
                ]

            stream_items = sorted(
                stream_items,
                key=lambda row: (
                    row.get("created_at").timestamp()
                    if getattr(row.get("created_at"), "timestamp", None)
                    else 0
                ),
                reverse=True,
            )
            paged_items, total = _paginate_list(stream_items, page=page, page_size=page_size)
            for row in paged_items:
                created_at = row.get("created_at")
                if hasattr(created_at, "isoformat"):
                    row["created_at"] = created_at.isoformat()
            return ok({"items": paged_items, "total": total, "page": page, "page_size": page_size})

        return ok({"plans": plans, "reviews": reviews, "submissions": submissions})


class TeacherMetricsStreamView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        aggregate_response = TeacherAggregateDashboardView().get(request)
        radar_response = TeacherClassRadarView().get(request)
        aggregate_payload = (aggregate_response.data or {}).get("data") or {}
        radar_payload = (radar_response.data or {}).get("data") or {}

        items = []
        for key, value in (aggregate_payload.get("summary") or {}).items():
            items.append({"key": key, "value": value, "source": "aggregate"})

        for item in (aggregate_payload.get("class_radar") or []):
            key = str(item.get("name") or item.get("rubric_id") or "").strip()
            if not key:
                continue
            items.append({"key": key, "value": item.get("score"), "source": "aggregate"})

        if isinstance(radar_payload, dict):
            for key, value in radar_payload.items():
                items.append({"key": key, "value": value, "source": "radar"})

        source = str(request.query_params.get("status") or request.query_params.get("source") or "").strip().lower()
        keyword = str(request.query_params.get("q") or "").strip().lower()
        if source in {"aggregate", "radar"}:
            items = [item for item in items if item.get("source") == source]
        if keyword:
            items = [
                item
                for item in items
                if keyword in str(item.get("key") or "").lower() or keyword in str(item.get("value") or "").lower()
            ]

        items = sorted(items, key=lambda row: (str(row.get("source") or ""), str(row.get("key") or "")))
        page_enabled, page, page_size = _read_paging(request, default_page_size=10, max_page_size=100)
        if page_enabled:
            paged_items, total = _paginate_list(items, page=page, page_size=page_size)
            return ok({"items": paged_items, "total": total, "page": page, "page_size": page_size})
        return ok(items)


class ReviewUpdateView(APIView):
    permission_classes = [IsTeacher]

    def put(self, request, review_id):
        review = get_object_or_404(ReviewRecord, id=review_id)
        if not _teacher_can_access_plan(request.user, review.plan, allow_plan_fallback=True):
            return fail("鏃犳潈闄愮紪杈戣鎵归槄", status.HTTP_403_FORBIDDEN)

        serializer = ReviewRecordSerializer(instance=review, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save(reviewer=request.user, audience_role=ReviewRecord.ROLE_TEACHER)
        return ok(serializer.data, message="审阅建议已更新")


class TeacherDashboardView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        student_ids = _teacher_student_ids(request.user, include_pending=False, allow_plan_fallback=True)
        plans = list(Plan.objects.filter(student_id__in=student_ids, status=Plan.STATUS_SUBMITTED))
        existing_review_count = (
            ReviewRecord.objects.filter(plan__in=plans, audience_role=ReviewRecord.ROLE_STUDENT).count()
            if plans
            else 0
        )
        rubrics = [
            {"code": rubric_id, "label": label, "weight": round(5.0 / max(len(RUBRIC_NAME_MAP), 1), 2)}
            for rubric_id, label in RUBRIC_NAME_MAP.items()
        ]

        dimension_scores = {item["code"]: [] for item in rubrics}
        for plan in plans:
            review_payload = _get_or_create_review_payload(plan, ReviewRecord.ROLE_STUDENT, reviewer=request.user)
            for dim in _normalize_review_dimension_scores(review_payload.get("dimension_scores") or []):
                rubric_id = str(dim.get("code") or "").upper()
                if rubric_id in dimension_scores:
                    dimension_scores[rubric_id].append(float(dim.get("score") or 0))

        avg_map = {
            code: (sum(values) / len(values) if values else 3.0)
            for code, values in dimension_scores.items()
        }

        plan_count = len(plans)
        submitted_student_count = len({plan.student_id for plan in plans})
        total_student_count = len(set(student_ids))
        metrics = {
            "student_count": total_student_count,
            "plan_count": plan_count,
            "submission_rate": round((submitted_student_count / total_student_count) * 100, 2) if total_student_count else 0,
            "optimization_rate": round((existing_review_count / plan_count) * 100, 2) if plan_count else 0,
            "case_library_count": CaseLibraryDocument.objects.filter(clean_status=CaseLibraryDocument.CLEAN_DONE).count(),
            "synthetic_case_count": CaseLibraryDocument.objects.filter(is_synthetic=True).count(),
            "updated_at": timezone.now(),
        }

        chart = {
            "indicators": [item["label"] for item in rubrics],
            "avg_scores": [round(avg_map.get(item["code"], 3.0), 2) for item in rubrics],
            "submission_trend": [
                {"name": "W1", "value": sum(1 for plan in plans if getattr(plan.created_at, "isoweekday", lambda: 0)() == 1)},
                {"name": "W2", "value": sum(1 for plan in plans if getattr(plan.created_at, "isoweekday", lambda: 0)() == 2)},
                {"name": "W3", "value": sum(1 for plan in plans if getattr(plan.created_at, "isoweekday", lambda: 0)() == 3)},
                {"name": "W4", "value": sum(1 for plan in plans if getattr(plan.created_at, "isoweekday", lambda: 0)() == 4)},
            ],
            "weights": [{"label": item["label"], "weight": item["weight"]} for item in rubrics],
        }
        return ok({"metrics": metrics, "chart": chart})


class CaseLibrarySummaryView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        queryset = CaseLibraryDocument.objects.filter(clean_status=CaseLibraryDocument.CLEAN_DONE)
        industry_stats = {}
        for item in queryset.values("industry_category"):
            key = item.get("industry_category") or "鍏朵粬"
            industry_stats[key] = industry_stats.get(key, 0) + 1

        latest = queryset.order_by("-updated_at")[:10]
        return ok(
            {
                "total": queryset.count(),
                "synthetic": queryset.filter(is_synthetic=True).count(),
                "industry_stats": industry_stats,
                "latest_items": CaseLibraryDocumentSerializer(latest, many=True).data,
            }
        )


class PromptSceneConfigListView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        queryset = PromptSceneConfig.objects.filter(is_active=True).order_by("scene_key")
        keyword = str(request.query_params.get("q") or "").strip()
        if keyword:
            queryset = queryset.filter(Q(scene_name__icontains=keyword) | Q(scene_key__icontains=keyword))

        page_enabled, page, page_size = _read_paging(request, default_page_size=8, max_page_size=50)
        if page_enabled:
            page_queryset, total = _paginate_queryset(queryset, page=page, page_size=page_size)
            data = PromptSceneConfigSerializer(page_queryset, many=True).data
            return ok({"items": data, "total": total, "page": page, "page_size": page_size})
        return ok(PromptSceneConfigSerializer(queryset, many=True).data)


class ScoringRubricListView(APIView):
    def get(self, request):
        queryset = ScoringRubric.objects.filter(is_active=True)
        return ok(ScoringRubricSerializer(queryset, many=True).data)


class TeacherCommonIssuesView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        student_ids = _teacher_student_ids(request.user, include_pending=False, allow_plan_fallback=True)
        plan_ids_raw = (request.query_params.get("plan_ids") or "").strip()
        selected_plan_ids = []
        if plan_ids_raw:
            selected_plan_ids = [int(item) for item in plan_ids_raw.split(",") if item.strip().isdigit()]
        if selected_plan_ids:
            plans = list(
                Plan.objects.filter(
                    id__in=selected_plan_ids,
                    student_id__in=student_ids,
                    status=Plan.STATUS_SUBMITTED,
                )
            )
            parsed_payloads = []
            for plan in plans:
                text = (plan.extracted_text or "").strip()
                if not text:
                    try:
                        pages = _ensure_plan_pages(plan)
                        text = "\n".join((item or {}).get("text", "") for item in pages).strip()
                    except Exception:
                        text = ""
                if text:
                    parsed_payloads.append(
                        {
                            "plan_id": plan.id,
                            "title": plan.title,
                            "student": plan.student.display_name,
                            "text": text,
                        }
                    )

            if parsed_payloads:
                aggregated = summarize_common_issues_from_plan_texts(parsed_payloads)
                if not aggregated:
                    # Fallback to historical review records when no reliable extraction is available.
                    reviews_qs = ReviewRecord.objects.filter(plan_id__in=[item["plan_id"] for item in parsed_payloads])
                    aggregated = summarize_common_issues(list(reviews_qs.values("issues")))
                if aggregated is not None:
                    CommonIssueReport.objects.filter(teacher=request.user).delete()
                    for item in aggregated:
                        CommonIssueReport.objects.create(teacher=request.user, **item)

        queryset = CommonIssueReport.objects.filter(teacher=request.user).order_by("-frequency", "-created_at")
        keyword = str(request.query_params.get("q") or "").strip()
        if keyword:
            queryset = queryset.filter(Q(problem_type__icontains=keyword) | Q(sample_case__icontains=keyword))

        page_enabled, page, page_size = _read_paging(request, default_page_size=8, max_page_size=50)
        if page_enabled:
            page_queryset, total = _paginate_queryset(queryset, page=page, page_size=page_size)
            data = CommonIssueReportSerializer(page_queryset, many=True).data
            return ok({"items": data, "total": total, "page": page, "page_size": page_size})
        return ok(CommonIssueReportSerializer(queryset, many=True).data)


class TeacherKnowledgeRecommendationView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        issues = CommonIssueReport.objects.filter(teacher=request.user).order_by("-frequency")[:10]
        payload = [
            {"problem_type": issue.problem_type, "frequency": issue.frequency, "sample_case": issue.sample_case}
            for issue in issues
        ]
        return ok(recommend_knowledge(payload))


class TeacherAIChatView(APIView):
    permission_classes = [IsTeacher]

    def post(self, request):
        question = request.data.get("question", "").strip()
        context = request.data.get("context", {}) or {}
        if not question:
            return fail("问题不能为空")

        history = TeacherChatRecord.objects.filter(teacher=request.user).order_by("-created_at")[:8]
        context["history"] = TeacherChatRecordSerializer(history, many=True).data
        answer = teacher_chat_answer(question, context)

        record = TeacherChatRecord.objects.create(
            teacher=request.user,
            context_type=context.get("type", "prep"),
            question=question,
            answer=answer,
        )
        return ok(TeacherChatRecordSerializer(record).data)


def _can_access_plan(user, plan: Plan) -> bool:
    if user.role == User.ROLE_STUDENT:
        return plan.student_id == user.id
    if user.role == User.ROLE_TEACHER:
        return _teacher_can_access_plan(user, plan, allow_plan_fallback=True)
    return False


def _safe_review_item_text(item, *keys: str) -> str:
    if isinstance(item, dict):
        for key in keys:
            value = str(item.get(key) or "").strip()
            if value:
                return value
        return ""
    return str(item or "").strip()


def _safe_plan_hypergraph_context(plan: Plan | None):
    if not plan:
        return None
    try:
        return build_plan_diagnosis(plan)
    except Exception as exc:
        log_event(
            "hypergraph_diagnosis_degraded",
            {"plan_id": getattr(plan, "id", None), "reason": str(exc)[:220]},
        )
        return None


def _run_agent_async_task(task_id: int, user_id: int, agent: str, payload: dict):
    close_old_connections()
    task = ConversationRecord.objects.filter(id=task_id).first()
    user = User.objects.filter(id=user_id).first()
    if not task or not user:
        return

    task.metadata = {**(task.metadata or {}), "status": "running"}
    task.save(update_fields=["metadata"])

    try:
        result = _dispatch_agent_task(user=user, agent=agent, payload=payload)
        task.answer = "done"
        task.metadata = {**(task.metadata or {}), "status": "done", "result": result}
        task.save(update_fields=["answer", "metadata"])
    except Exception as exc:
        task.metadata = {**(task.metadata or {}), "status": "error", "error": str(exc)}
        task.save(update_fields=["metadata"])


def _dispatch_agent_task(user, agent: str, payload: dict):
    if agent == "guide":
        mode = str(payload.get("mode") or "").strip().lower()
        if mode in {"opening", "teach", "tutor", "knowledge", "practice"}:
            tutor = TutorAgent()
            if mode == "opening":
                return tutor.opening()
            return tutor.chat(str(payload.get("message") or ""), payload.get("history") or [])
        stage = str(payload.get("stage") or "idea")
        answer = str(payload.get("answer") or "")
        return {"next": next_guiding_question(stage, answer)}

    if agent == "tutor":
        mode = str(payload.get("mode") or "chat")
        tutor = TutorAgent()
        if mode == "opening":
            return tutor.opening()
        return tutor.chat(str(payload.get("message") or ""), payload.get("history") or [])

    if agent == "pitch_optimize":
        plan_id = payload.get("plan_id")
        plan = get_object_or_404(Plan, id=plan_id, student=user)
        merged = {
            **payload,
            "plan_text": plan.extracted_text or "",
        }
        result = PitchOptimizeAgent().chat(merged)
        if isinstance(result, dict) and result.get("ok") is False:
            raise RuntimeError(str(result.get("message") or "鍑轰簡浜涘皬闂"))
        return result

    if agent == "pitch_simulate":
        sim = PitchSimulationAgent()
        action = str(payload.get("action") or "start")
        if action == "start":
            result = sim.start(payload)
            if isinstance(result, dict) and result.get("ok") is False:
                raise RuntimeError(str(result.get("message") or "鍑轰簡浜涘皬闂"))
            return result
        if action == "answer":
            result = sim.ask_next(payload)
            if isinstance(result, dict) and result.get("ok") is False:
                raise RuntimeError(str(result.get("message") or "鍑轰簡浜涘皬闂"))
            return result
        if action == "end":
            result = sim.end(payload)
            if isinstance(result, dict) and result.get("ok") is False:
                raise RuntimeError(str(result.get("message") or "鍑轰簡浜涘皬闂"))
            return result
        if action == "qa":
            if not bool(payload.get("after_end")):
                raise RuntimeError("路演模拟未结束前不提供答疑")
            result = sim.qa_after_end(payload)
            if isinstance(result, dict) and result.get("ok") is False:
                raise RuntimeError(str(result.get("message") or "鍑轰簡浜涘皬闂"))
            return result
        raise RuntimeError("鏈煡妯℃嫙鍔ㄤ綔")

    if agent == "competition_coach":
        plan_id = payload.get("plan_id")
        plan = get_object_or_404(Plan, id=plan_id, student=user) if plan_id else None
        question = str(payload.get("question") or "").strip()
        history = payload.get("history") or []
        plan_text = str(payload.get("text") or "").strip()
        if plan and not plan_text:
            plan_text = plan.extracted_text or ""
        if not plan_text:
            plan_text = _history_to_context_text(history) or question
        hg_report = _safe_plan_hypergraph_context(plan)
        return CompetitionCoachAgent().chat(
            competition=str(payload.get("competition") or "挑战杯"),
            plan_text=plan_text,
            question=question,
            history=history,
            hypergraph_context=hg_report,
            project_type=str(payload.get("project_type") or "auto"),
        )

    if agent == "financial_design":
        plan_id = payload.get("plan_id")
        plan_text = str(payload.get("text") or "").strip()
        plan = None
        if plan_id and not plan_text:
            plan = get_object_or_404(Plan, id=plan_id, student=user)
            plan_text = plan.extracted_text or ""
        if not plan_text:
            raise RuntimeError("请先选择方案或输入项目内容")
        return FinancialDesignAgent().generate(
            plan_text=plan_text,
            question=str(payload.get("question") or ""),
            project_type=str(payload.get("project_type") or "auto"),
            competition=str(payload.get("competition") or ""),
            history=payload.get("history") or [],
            current_plan=payload.get("current_plan"),
            hypergraph_context=_safe_plan_hypergraph_context(plan),
        )

    if agent == "pdf_chat":
        plan_id = payload.get("plan_id")
        question = str(payload.get("question") or "").strip()
        if not question:
            raise RuntimeError("question涓嶈兘涓虹┖")
        plan = get_object_or_404(Plan, id=plan_id, student=user)
        review_id = payload.get("review_id")
        latest_review = ReviewRecord.objects.filter(id=review_id, plan=plan).first() if review_id else ReviewRecord.objects.filter(plan=plan, audience_role=ReviewRecord.ROLE_STUDENT).first()
        context = {
            "plan_title": plan.title,
            "plan_excerpt": (plan.extracted_text or "")[:2500],
            "review_summary": latest_review.summary if latest_review else "",
            "guidance_questions": latest_review.guidance_questions if latest_review else [],
            "annotations": latest_review.annotations if latest_review else [],
            "hypergraph": ((latest_review.review_meta if latest_review else {}) or {}).get("hypergraph") or build_plan_diagnosis(plan),
        }
        answer = student_pdf_chat_answer(question, context)
        return {"answer": answer, "context": context}

    raise RuntimeError("未知智能体类型")


class AgentChatAsyncSubmitView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        agent = str(request.data.get("agent") or "").strip()
        payload = request.data.get("payload") or {}
        if not agent:
            return fail("agent涓嶈兘涓虹┖")

        task = ConversationRecord.objects.create(
            user=request.user,
            stage="agent_async",
            question=f"agent={agent}",
            metadata={"status": "queued", "agent": agent, "payload": payload},
        )
        thread = threading.Thread(target=_run_agent_async_task, args=(task.id, request.user.id, agent, payload), daemon=True)
        thread.start()
        return ok({"task_id": task.id, "status": "queued"})


class AgentChatTaskStatusView(APIView):
    permission_classes = [IsStudent]

    def get(self, request, task_id):
        task = get_object_or_404(ConversationRecord, id=task_id, user=request.user, stage="agent_async")
        meta = task.metadata or {}
        return ok(
            {
                "task_id": task.id,
                "status": meta.get("status", "queued"),
                "result": meta.get("result"),
                "error": meta.get("error"),
            }
        )


class RoadmapSimulateAgentView(APIView):
    def post(self, request):
        simulator = RoadmapSimulator()
        if request.data.get("mode") == "report":
            turns = request.data.get("turns") or []
            return ok(simulator.build_report(turns))

        question_index = request.data.get("question_index")
        answer = request.data.get("answer", "")
        if question_index is None:
            return ok(simulator.start())
        return ok(simulator.next_turn(question_index, answer))


class CompetitionAdviseAgentView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        competition = (request.data.get("competition") or "挑战杯").strip()
        text = (request.data.get("text") or "").strip()
        project_type = str(request.data.get("project_type") or "auto").strip().lower()
        plan_id = request.data.get("plan_id")
        plan = None
        if plan_id:
            plan = get_object_or_404(Plan, id=plan_id)
            if not _can_access_plan(request.user, plan):
                return fail("鏃犳潈闄愯闂鏂规", status.HTTP_403_FORBIDDEN)
            if not text:
                text = plan.extracted_text or ""
        if not text:
            return fail("请先选择方案或输入项目摘要")

        advisor = CompetitionAdvisor()
        try:
            return ok(
                advisor.advise(
                    competition=competition,
                    text=text,
                    project_type=project_type,
                    hypergraph_context=build_plan_diagnosis(plan) if plan else None,
                )
            )
        except Exception:
            return fail("鍑轰簡浜涘皬闂")


class CompetitionCoachChatView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        competition = (request.data.get("competition") or "挑战杯").strip()
        mode = str(request.data.get("mode") or "chat").strip().lower()
        question = (request.data.get("question") or "").strip()
        history = request.data.get("history") or []
        plan_id = request.data.get("plan_id")
        plan_text = (request.data.get("text") or "").strip()
        project_type = str(request.data.get("project_type") or "auto").strip()

        plan = None
        if plan_id:
            plan = get_object_or_404(Plan, id=plan_id, student=request.user)
        if plan_id and not plan_text:
            plan_text = plan.extracted_text or ""
        if not plan_text:
            plan_text = _history_to_context_text(history) or question
        if mode == "chat" and not question:
            return fail("question涓嶈兘涓虹┖")
        if mode == "co_write" and not plan_text:
            return fail("请先输入协作上下文，或先发起几轮竞赛教练对话")

        coach = CompetitionCoachAgent()
        if mode == "co_write":
            result = coach.co_write_plan(
                competition=competition,
                plan_text=plan_text,
                history=history,
                project_type=project_type,
            )
            return ok(result)

        result = coach.chat(
            competition=competition,
            plan_text=plan_text,
            question=question,
            history=history,
            hypergraph_context=build_plan_diagnosis(plan) if plan_id else None,
            project_type=project_type,
        )
        return ok(result)


class CompetitionCoachExportWordView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        competition = (request.data.get("competition") or "挑战杯").strip()
        plan_id = request.data.get("plan_id")
        plan_text = (request.data.get("text") or "").strip()
        history = request.data.get("history") or []
        project_type = str(request.data.get("project_type") or "auto").strip()

        if plan_id and not plan_text:
            plan = get_object_or_404(Plan, id=plan_id, student=request.user)
            plan_text = plan.extracted_text or ""
        if not plan_text:
            plan_text = _history_to_context_text(history)
        if not plan_text:
            return fail("请先与竞赛教练协作生成内容，再导出 Word")

        result = CompetitionCoachAgent().co_write_plan(
            competition=competition,
            plan_text=plan_text,
            history=history,
            project_type=project_type,
        )
        markdown_text = str(result.get("markdown") or "").strip()
        if not markdown_text:
            sections = result.get("sections") or []
            lines = [f"# {result.get('title') or '竞赛计划书'}", ""]
            for item in sections:
                lines.append(f"## {item.get('name') or '绔犺妭'}")
                lines.append(str(item.get("content") or "").strip())
                lines.append("")
            markdown_text = "\n".join(lines).strip()

        content = _markdown_to_docx_bytes(result.get("title") or f"{competition}计划书", markdown_text)
        filename = f"{competition}_计划书.docx"
        response = HttpResponse(
            content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response


class FinancialDesignAgentView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        plan_id = request.data.get("plan_id")
        plan_text = (request.data.get("text") or "").strip()
        plan = None
        if plan_id and not plan_text:
            plan = get_object_or_404(Plan, id=plan_id, student=request.user)
            plan_text = plan.extracted_text or ""
        if not plan_text:
            return fail("请先选择方案或输入项目摘要")

        hypergraph_context = _safe_plan_hypergraph_context(plan)
        try:
            payload = FinancialDesignAgent().generate(
                plan_text=plan_text,
                question=str(request.data.get("question") or ""),
                project_type=str(request.data.get("project_type") or "auto"),
                competition=str(request.data.get("competition") or ""),
                history=request.data.get("history") or [],
                current_plan=request.data.get("current_plan"),
                hypergraph_context=hypergraph_context,
            )
        except Exception as exc:
            log_event(
                "financial_design_generate_failed",
                {"plan_id": getattr(plan, "id", None), "reason": str(exc)[:220]},
            )
            # Keep the endpoint available even when external context/parsing fails.
            payload = FinancialDesignAgent().generate(
                plan_text=plan_text,
                question=str(request.data.get("question") or ""),
                project_type=str(request.data.get("project_type") or "auto"),
                competition=str(request.data.get("competition") or ""),
                history=request.data.get("history") or [],
                current_plan=request.data.get("current_plan"),
                hypergraph_context=None,
            )
        return ok(payload)


class ProjectBottleneckView(APIView):
    def get(self, request, plan_id):
        plan = get_object_or_404(Plan, id=plan_id)
        if not _can_access_plan(request.user, plan):
            return fail("鏃犳潈闄愯闂鏂规", status.HTTP_403_FORBIDDEN)

        quality = evaluate_project_quality(plan)
        coach = ProjectCoach()
        diagnosis = coach.diagnose(quality.get("facts") or {})
        return ok(diagnosis)


class TutorAgentView(APIView):
    def post(self, request):
        if request.data.get("mode") == "opening":
            return ok(TutorAgent().opening())

        if request.data.get("mode") == "chat":
            message = (request.data.get("message") or "").strip()
            history = request.data.get("history") or []
            return ok(TutorAgent().chat(message=message, history=history))

        topic = (request.data.get("topic") or "").strip()
        if not topic:
            return fail("topic涓嶈兘涓虹┖")
        tutor = TutorAgent()
        return ok(tutor.teach(topic))


class PitchOptimizeAgentView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        plan_id = request.data.get("plan_id")
        if not plan_id:
            return fail("浼樺寲鏅鸿兘浣撳繀椤诲厛閫夋嫨宸叉彁浜ょ殑鏂规(plan_id)")
        plan = get_object_or_404(Plan, id=plan_id, student=request.user)

        ppt_file = request.FILES.get("ppt_file")
        script_file = request.FILES.get("script_file")
        script_text = request.data.get("script_text") or ""
        if script_file and not script_text:
            try:
                script_text = script_file.read().decode("utf-8", errors="ignore")[:3000]
            except Exception:
                script_text = ""

        payload = {
            "plan_text": plan.extracted_text or "",
            "message": request.data.get("message") or "",
            "duration_minutes": request.data.get("duration_minutes") or 6,
            "ppt_name": request.data.get("ppt_name") or (ppt_file.name if ppt_file else ""),
            "script_name": request.data.get("script_name") or (script_file.name if script_file else ""),
            "script_text": script_text,
        }
        return ok(PitchOptimizeAgent().chat(payload))


class PitchSimulationAgentView(APIView):
    permission_classes = [IsStudent]

    def post(self, request):
        action = (request.data.get("action") or "").strip() or "start"
        agent = PitchSimulationAgent()
        payload = dict(request.data or {})
        if request.FILES.get("ppt_file"):
            payload["ppt_name"] = request.FILES["ppt_file"].name
        if request.FILES.get("script_file"):
            payload["script_file_name"] = request.FILES["script_file"].name

        if action == "start":
            return ok(agent.start(payload))
        if action == "answer":
            return ok(agent.ask_next(payload))
        if action == "end":
            return ok(agent.end(payload))
        if action == "qa":
            if not bool(payload.get("after_end")):
                return fail("路演模拟未结束前不提供答疑，请先结束模拟。")
            return ok(agent.qa_after_end(payload))
        return fail("action仅支持start/answer/end/qa")


class TeacherClassRadarView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        student_ids = _teacher_student_ids(request.user, include_pending=False, allow_plan_fallback=True)
        plans = list(Plan.objects.filter(student_id__in=student_ids, status=Plan.STATUS_SUBMITTED))
        if not plans:
            return ok({})

        buckets = defaultdict(list)
        for plan in plans:
            review_payload = _get_or_create_review_payload(plan, ReviewRecord.ROLE_STUDENT, reviewer=request.user)
            meta = review_payload.get("review_meta") or {}
            if not bool(meta.get("is_complete", _is_plan_complete(plan))):
                continue
            for dim in _normalize_review_dimension_scores(review_payload.get("dimension_scores") or []):
                rubric_id = str(dim.get("code") or "").upper()
                if rubric_id in RUBRIC_NAME_MAP:
                    buckets[rubric_id].append(float(dim.get("score") or 0))
        return ok({rubric_id: round(sum(values) / len(values), 2) for rubric_id, values in buckets.items()})


class TeacherCommonMistakesView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        student_ids = _teacher_student_ids(request.user, include_pending=False, allow_plan_fallback=True)
        plans = Plan.objects.filter(student_id__in=student_ids, status=Plan.STATUS_SUBMITTED)

        counter = {}
        for plan in plans[:50]:
            quality = evaluate_project_quality(plan)
            for item in quality.get("triggered_rules") or []:
                rid = item.get("rule_id")
                if not rid:
                    continue
                row = counter.setdefault(rid, {"rule_id": rid, "description": item.get("description", ""), "count": 0, "project_ids": []})
                row["count"] += 1
                row["project_ids"].append(plan.id)

        top5 = sorted(counter.values(), key=lambda x: x["count"], reverse=True)[:5]
        return ok(top5)


class TeacherHighRiskProjectsView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        student_ids = _teacher_student_ids(request.user, include_pending=False, allow_plan_fallback=True)
        plans = Plan.objects.filter(student_id__in=student_ids, status=Plan.STATUS_SUBMITTED).select_related("student")
        risks = []
        for plan in plans[:100]:
            quality = evaluate_project_quality(plan)
            high_rules = [r for r in (quality.get("triggered_rules") or []) if r.get("severity") == "high"]
            total = float(quality.get("total_score_100") or 0)
            if high_rules or total < 60:
                risk_level = "high" if high_rules else "medium"
                risks.append(
                    {
                        "project_id": plan.id,
                        "project_name": plan.title,
                        "student_name": plan.student.display_name,
                        "risk_types": [r.get("rule_id") for r in high_rules] or ["low_rubric_score"],
                        "total_score": total,
                        "risk_level": risk_level,
                        "intervention": "优先安排一对一答辩演练，并补齐关键证据链。",
                    }
                )
        keyword = str(request.query_params.get("q") or "").strip().lower()
        status_filter = str(request.query_params.get("status") or "").strip().lower()
        if keyword:
            risks = [
                item
                for item in risks
                if keyword in str(item.get("project_name") or "").lower()
                or keyword in str(item.get("student_name") or "").lower()
            ]
        if status_filter in {"high", "medium", "low"}:
            risks = [item for item in risks if str(item.get("risk_level") or "").lower() == status_filter]

        page_enabled, page, page_size = _read_paging(request, default_page_size=8, max_page_size=50)
        if page_enabled:
            paged_items, total = _paginate_list(risks, page=page, page_size=page_size)
            return ok({"items": paged_items, "total": total, "page": page, "page_size": page_size})
        return ok(risks)


class ProjectRubricScoreDetailView(APIView):
    def get(self, request, plan_id):
        plan = get_object_or_404(Plan, id=plan_id)
        if not _can_access_plan(request.user, plan):
            return fail("鏃犳潈闄愯闂鏂规", status.HTTP_403_FORBIDDEN)

        quality = evaluate_project_quality(plan) if (plan.extracted_text or "").strip() else {"triggered_rules": []}
        items = ProjectRubricScore.objects.filter(project=plan).select_related("rubric_item")
        serialized = ProjectRubricScoreSerializer(items, many=True).data
        total = 0.0
        ability_profile = []
        for row in items:
            total += float(row.score) * float(row.rubric_item.weight)
            rid = row.rubric_item.rubric_id.upper()
            ability_profile.append({
                "rubric_id": rid,
                "name": RUBRIC_NAME_MAP.get(rid, row.rubric_item.name),
                "score": round(float(row.score), 2),
            })

        triggered = quality.get("triggered_rules") or []
        top_rule = None
        if triggered:
            top_rule = sorted(
                triggered,
                key=lambda item: (-SEVERITY_ORDER.get(str(item.get("severity") or "").lower(), 0), str(item.get("rule_id") or "")),
            )[0]

        return ok(
            {
                "project_id": plan.id,
                "total_score_100": round(total * 10, 2),
                "items": serialized,
                "hypergraph_rubric_basis": quality.get("hypergraph_rubric_basis") or {},
                "ability_profile": ability_profile,
                "top_risk_rule": top_rule,
                "improvement_tip": (top_rule or {}).get("fix_task") or "建议优先补齐关键证据，再迭代商业闭环。",
            }
        )


class TeacherAggregateDashboardView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        student_ids = _teacher_student_ids(request.user, include_pending=False, allow_plan_fallback=True)
        student_count = len(set(student_ids))
        all_submitted = list(
            Plan.objects.filter(student_id__in=student_ids, status=Plan.STATUS_SUBMITTED)
            .select_related("student")
            .order_by("student_id", "title", "-version", "-created_at")
        )
        submitted_student_count = len({plan.student_id for plan in all_submitted})
        submission_rate = _safe_ratio(submitted_student_count, student_count) if student_count else 0.0
        if not all_submitted:
            return ok(
                {
                    "selector": {"projects": []},
                    "complete": {"project_count": 0, "avg_rubric_score": 0.0, "class_radar": [], "top_issues": []},
                    "incomplete": {"project_count": 0, "avg_progress": 0.0, "missing_top5": []},
                    "plagiarism_check": {
                        "enabled": True,
                        "checked_plan_count": 0,
                        "suspicious_pair_count": 0,
                        "high_risk_count": 0,
                        "suspicious_pairs": [],
                    },
                    "process_evaluation": {"overview": {"avg_process_score": 0.0, "project_count": 0, "high_process_count": 0, "low_process_count": 0}, "project_rows": []},
                    "class_warning": {"warning_count": 0, "items": []},
                    "metrics": {
                        "student_count": student_count,
                        "submitted_student_count": submitted_student_count,
                        "plan_count": 0,
                        "submission_rate": round(submission_rate, 2),
                        "avg_rubric_score": 0.0,
                        "plagiarism_suspect_pairs": 0,
                        "warning_count": 0,
                        "low_process_project_count": 0,
                    },
                    "summary": {"class_avg_score": 0.0, "high_risk_project_count": 0, "weakest_dimension": "-"},
                    "class_radar": [],
                    "rule_heatmap": [],
                }
            )

        selected_raw = (request.query_params.get("selected_plan_ids") or "").strip()
        selected_ids = {int(x) for x in selected_raw.split(",") if x.strip().isdigit()} if selected_raw else set()

        grouped = defaultdict(list)
        for plan in all_submitted:
            grouped[_to_project_group_key(plan)].append(plan)

        selected_plans = []
        selector_projects = []
        for key, versions in grouped.items():
            versions = sorted(versions, key=lambda p: (p.version, p.created_at), reverse=True)
            default_plan = versions[0]
            picked = None
            if selected_ids:
                picked = next((p for p in versions if p.id in selected_ids), None)
            if selected_ids:
                if picked:
                    selected_plans.append(picked)
            else:
                selected_plans.append(default_plan)

            selector_projects.append(
                {
                    "project_key": key,
                    "title": default_plan.title,
                    "student_id": default_plan.student_id,
                    "student_name": default_plan.student.display_name,
                    "default_plan_id": default_plan.id,
                    "selected_plan_id": picked.id if picked else (default_plan.id if not selected_ids else None),
                    "is_selected": bool(picked) if selected_ids else True,
                    "versions": [
                        {
                            "plan_id": p.id,
                            "version": p.version,
                            "created_at": p.created_at,
                            "is_default": p.id == default_plan.id,
                            "is_complete": _is_plan_complete(p),
                            "progress": _plan_progress(p),
                        }
                        for p in versions
                    ],
                }
            )

        review_by_plan = {}
        rubric_by_plan = {}
        complete_plans = []
        incomplete_plans = []
        for plan in selected_plans:
            review_payload = _get_or_create_review_payload(plan, ReviewRecord.ROLE_STUDENT, reviewer=request.user)
            review_meta = (review_payload.get("review_meta") or {}).copy()
            review_by_plan[plan.id] = review_payload

            dimension_scores = _normalize_review_dimension_scores(
                review_payload.get("dimension_scores") or review_meta.get("dimension_scores") or []
            )
            rubric_by_plan[plan.id] = {"dimensions": dimension_scores}

            is_complete = bool((review_payload.get("review_meta") or {}).get("is_complete", _is_plan_complete(plan)))
            if is_complete:
                complete_plans.append(plan)
            else:
                incomplete_plans.append(plan)

        rubric_bucket = defaultdict(list)
        all_project_scores = []
        plan_avg_score_map = {}
        rubric_name_to_id = {v: k for k, v in RUBRIC_NAME_MAP.items()}
        for plan in selected_plans:
            generated = review_by_plan.get(plan.id) or {}
            rubric_payload = rubric_by_plan.get(plan.id) or {}
            dim_rows = rubric_payload.get("dimensions") or []

            if dim_rows:
                # 1-5 scale average.
                avg_score = sum(float(x.get("score") or 0) for x in dim_rows) / max(len(dim_rows), 1)
                rounded_score = round(avg_score, 2)
                all_project_scores.append(rounded_score)
                plan_avg_score_map[plan.id] = rounded_score
            for row in dim_rows:
                label = str(row.get("label") or row.get("dimension") or "").strip()
                rid = rubric_name_to_id.get(label) or str(row.get("code") or "").upper()
                if rid in RUBRIC_NAME_MAP:
                    rubric_bucket[rid].append(float(row.get("score") or 0))

        class_radar = []
        for rid in [f"R{i}" for i in range(1, 10)]:
            vals = rubric_bucket.get(rid, [])
            class_radar.append({"rubric_id": rid, "name": RUBRIC_NAME_MAP.get(rid, rid), "score": round(sum(vals) / len(vals), 2) if vals else 0.0})

        top_issues = _build_common_issue_rows(complete_plans, review_by_plan)

        missing_counter = defaultdict(lambda: {"count": 0, "urgency": "medium", "projects": []})
        progress_values = []
        for plan in incomplete_plans:
            generated = review_by_plan.get(plan.id) or {}
            meta = (generated.get("review_meta") or {})
            progress = int(meta.get("progress") or _plan_progress(plan) or 0)
            progress_values.append(progress)
            for miss in (meta.get("missing_parts") or []):
                part = str(miss.get("part") or "待补充部分").strip()
                if not part:
                    continue
                urgency = str(miss.get("urgency") or "medium").lower()
                row = missing_counter[part]
                row["count"] += 1
                row["urgency"] = urgency if urgency in {"low", "medium", "high"} else "medium"
                row["projects"].append({
                    "project_id": plan.id,
                    "project_name": plan.title,
                    "student_name": plan.student.display_name,
                    "answer": generated.get("summary") or "",
                })

        missing_top5 = [
            {"part": k, "count": v["count"], "urgency": v["urgency"], "projects": v["projects"]}
            for k, v in sorted(missing_counter.items(), key=lambda item: item[1]["count"], reverse=True)[:5]
        ]
        complete_project_scores = [plan_avg_score_map[plan.id] for plan in complete_plans if plan.id in plan_avg_score_map]
        plagiarism_check = _build_plagiarism_check(selected_plans)
        process_evaluation = _build_process_evaluation(selected_plans, grouped, review_by_plan, request.user)
        class_warning = _build_class_warning_payload(
            selected_plans=selected_plans,
            incomplete_plans=incomplete_plans,
            review_by_plan=review_by_plan,
            plagiarism_check=plagiarism_check,
            process_evaluation=process_evaluation,
        )

        return ok(
            {
                "selector": {"projects": selector_projects},
                "complete": {
                    "project_count": len(complete_plans),
                    "avg_rubric_score": round(sum(complete_project_scores) / len(complete_project_scores), 2) if complete_project_scores else 0.0,
                    "class_radar": class_radar,
                    "top_issues": top_issues,
                },
                "incomplete": {
                    "project_count": len(incomplete_plans),
                    "avg_progress": round(sum(progress_values) / len(progress_values), 2) if progress_values else 0.0,
                    "missing_top5": missing_top5,
                },
                "plagiarism_check": plagiarism_check,
                "process_evaluation": process_evaluation,
                "class_warning": class_warning,
                # Backward-compatible keys for existing tests/pages.
                "metrics": {
                    "student_count": student_count,
                    "submitted_student_count": submitted_student_count,
                    "plan_count": len(all_submitted),
                    "submission_rate": round(submission_rate, 2),
                    "avg_rubric_score": round(sum(all_project_scores) / len(all_project_scores), 2) if all_project_scores else 0.0,
                    "plagiarism_suspect_pairs": int(plagiarism_check.get("suspicious_pair_count") or 0),
                    "warning_count": int(class_warning.get("warning_count") or 0),
                    "low_process_project_count": int((process_evaluation.get("overview") or {}).get("low_process_count") or 0),
                },
                "summary": {
                    "class_avg_score": round(sum(all_project_scores) / len(all_project_scores), 2) if all_project_scores else 0.0,
                    "high_risk_project_count": len(incomplete_plans),
                    "weakest_dimension": min(class_radar, key=lambda x: x.get("score", 0)).get("name", "-") if class_radar else "-",
                },
                "class_radar": class_radar,
                "rule_heatmap": [],
            }
        )


class TeacherRuleDrilldownView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        rule_id = str(request.query_params.get("rule_id") or "").strip().upper()
        if not rule_id:
            return fail("rule_id涓嶈兘涓虹┖")

        student_ids = _teacher_student_ids(request.user, include_pending=False, allow_plan_fallback=True)
        plans = list(Plan.objects.filter(student_id__in=student_ids, status=Plan.STATUS_SUBMITTED).select_related("student"))
        items = []

        for plan in plans[:300]:
            quality = evaluate_project_quality(plan)
            hit = next((r for r in (quality.get("triggered_rules") or []) if str(r.get("rule_id") or "").upper() == rule_id), None)
            if not hit:
                continue

            evidence_qs = EvidenceRecord.objects.filter(
                project=plan,
                source=EvidenceRecord.SOURCE_MANUAL,
            ).filter(Q(extra__kind="rule") & Q(extra__rule_id=rule_id))[:3]

            chats = ConversationRecord.objects.filter(user=plan.student, metadata__plan_id=plan.id).order_by("created_at")[:2]
            chat_excerpt = [x.answer for x in chats if (x.answer or "").strip()]

            items.append(
                {
                    "project_id": plan.id,
                    "project_name": plan.title,
                    "student_name": plan.student.display_name,
                    "trigger_reason": hit.get("trigger_message") or "",
                    "evidence": [e.quote for e in evidence_qs] or [hit.get("trigger_message") or ""],
                    "ai_feedback_excerpt": chat_excerpt,
                }
            )

        return ok({"rule_id": rule_id, "projects": items})


class TeacherProjectEvidenceDetailView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request, plan_id):
        plan = get_object_or_404(Plan, id=plan_id)
        if not _teacher_can_access_plan(request.user, plan, allow_plan_fallback=True):
            return fail("鏃犳潈闄愯闂椤圭洰", status.HTTP_403_FORBIDDEN)
        if plan.status != Plan.STATUS_SUBMITTED:
            return fail("浠呮敮鎸佹煡鐪嬪凡鎻愪氦缁欐暀甯堢殑椤圭洰", status.HTTP_400_BAD_REQUEST)

        generated = _get_or_create_review_payload(plan, ReviewRecord.ROLE_STUDENT, reviewer=request.user)
        issues = (generated.get("issues") or [])
        annotations = (generated.get("annotations") or [])
        meta = (generated.get("review_meta") or {})
        chats = ConversationRecord.objects.filter(user=plan.student, metadata__plan_id=plan.id).order_by("created_at")[:2]
        chat_excerpt = [x.answer for x in chats if (x.answer or "").strip()]

        return ok(
            {
                "project_id": plan.id,
                "project_name": plan.title,
                "rules": [
                    {
                        "rule_id": str(idx + 1),
                        "severity": item.get("risk_level", "medium"),
                        "trigger_reason": item.get("description", ""),
                        "evidence": [item.get("snippet", "")],
                    }
                    for idx, item in enumerate(issues)
                ],
                "is_complete": bool(meta.get("is_complete", True)),
                "progress": int(meta.get("progress") or 0),
                "missing_parts": meta.get("missing_parts") or [],
                "issues": issues,
                "annotations": annotations,
                "summary": generated.get("summary") or "",
                "hypergraph": meta.get("hypergraph") or {"provenance": {"node_sources": [], "hyperedge_evidence": []}},
                "ai_feedback_excerpt": chat_excerpt,
            }
        )


class ProjectEvidenceChainView(APIView):
    def get(self, request, plan_id):
        plan = get_object_or_404(Plan, id=plan_id)
        if not _can_access_plan(request.user, plan):
            return fail("鏃犳潈闄愯闂鏂规", status.HTTP_403_FORBIDDEN)
        queryset = EvidenceRecord.objects.filter(project=plan)
        return ok(EvidenceRecordSerializer(queryset, many=True).data)


class TeacherInterventionCreateView(APIView):
    permission_classes = [IsTeacher]

    def post(self, request):
        plan_id = request.data.get("project") or request.data.get("plan_id")
        if not plan_id:
            return fail("project鎴杙lan_id涓嶈兘涓虹┖")
        plan = get_object_or_404(Plan, id=plan_id)
        if not _teacher_can_access_plan(request.user, plan, allow_plan_fallback=True):
            return fail("鏃犳潈闄愬共棰勮鏂规", status.HTTP_403_FORBIDDEN)

        TeacherIntervention.objects.filter(teacher=request.user, project=plan, is_active=True).update(is_active=False)
        payload = dict(request.data)
        payload["project"] = plan.id
        payload.setdefault("forced_question_points", [])
        payload.setdefault("override_rules", [])
        serializer = TeacherInterventionSerializer(data=payload)
        serializer.is_valid(raise_exception=True)
        serializer.save(teacher=request.user, project=plan, is_active=True)
        return ok(serializer.data, status_code=status.HTTP_201_CREATED)


class HypergraphImportCaseView(APIView):
    permission_classes = [IsTeacher]

    def post(self, request):
        payload = request.data or {}
        client = get_neo4j_client()
        result = client.import_case(payload)

        hyperedges = payload.get("hyperedges") or []
        edge_results = []
        for edge in hyperedges:
            edge_results.append(client.create_hyperedge(edge))

        return ok({"case": result, "hyperedges": edge_results, "neo4j_enabled": client.enabled})


class ProjectPotentialReportView(APIView):
    def get(self, request, plan_id):
        plan = get_object_or_404(Plan, id=plan_id)
        if not _can_access_plan(request.user, plan):
            return fail("鏃犳潈闄愯闂鏂规", status.HTTP_403_FORBIDDEN)
        return ok(build_potential_report(plan))


class TeacherClassLearningReportView(APIView):
    permission_classes = [IsTeacher]

    def get(self, request):
        return ok(build_class_learning_report(request.user))


class CaseRecallView(APIView):
    def post(self, request):
        query = (request.data.get("query") or "").strip()
        if not query:
            return fail("query涓嶈兘涓虹┖")
        top_k = int(request.data.get("top_k") or 5)
        industry = (request.data.get("industry") or "").strip()
        payload = VectorRetrievalService().recall_cases_with_meta(query=query, top_k=top_k, industry=industry)
        return ok(payload)


class WorkflowOrchestrateView(APIView):
    def post(self, request):
        payload = request.data or {}
        result = LangGraphWorkflowAgent().run(payload)
        return ok(result)


class ProjectHypergraphReasoningView(APIView):
    def get(self, request, plan_id):
        plan = get_object_or_404(Plan, id=plan_id)
        if not _can_access_plan(request.user, plan):
            return fail("鏃犳潈闄愯闂鏂规", status.HTTP_403_FORBIDDEN)
        return ok(infer_project_paths(plan))


class HypergraphPreviewGraphView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        limit = int(request.query_params.get("limit") or 80)
        return ok(build_preview_graph(limit=limit))


